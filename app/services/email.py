from flask import render_template, url_for, current_app
from flask_mail import Message
from ..utils import config_or_setting, generate_stage_ics, carried_amendment_summary
from ..utils import markdown_to_html

from ..extensions import mail, db
from ..models import (
    Member,
    Meeting,
    UnsubscribeToken,
    AppSetting,
    EmailLog,
    EmailSetting,
    Amendment,
    AmendmentObjection,
    User,
    Role,
    Motion,
    MotionSubmission,
    AmendmentSubmission,
    SubmissionToken,
)
from uuid6 import uuid7
from sqlalchemy import func
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os

DEFAULT_EMAIL_WHY_TEXT = (
    "You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails"
)


def auto_send_enabled(meeting: Meeting, kind: str) -> bool:
    """Return True if automatic emails of given type are enabled for meeting."""
    if AppSetting.get("manual_email_mode") == "1":
        return False
    setting = EmailSetting.query.filter_by(meeting_id=meeting.id, email_type=kind).first()
    return setting.auto_send if setting else True


def _log_email(member: Member, meeting: Meeting, kind: str, test_mode: bool) -> None:
    entry = EmailLog(
        meeting_id=meeting.id if meeting else None,
        member_id=member.id if member else None,
        type=kind,
        is_test=test_mode,
    )
    db.session.add(entry)
    db.session.commit()


def _unsubscribe_url(member: Member) -> str:
    token = UnsubscribeToken.query.filter_by(member_id=member.id).first()
    if not token:
        token = UnsubscribeToken(token=str(uuid7()), member_id=member.id)
        db.session.add(token)
        db.session.commit()
    return url_for('notifications.unsubscribe', token=token.token, _external=True)


def _resubscribe_url(member: Member) -> str:
    token = UnsubscribeToken.query.filter_by(member_id=member.id).first()
    if not token:
        token = UnsubscribeToken(token=str(uuid7()), member_id=member.id)
        db.session.add(token)
        db.session.commit()
    return url_for('notifications.resubscribe', token=token.token, _external=True)


def _sender() -> str:
    """Return configured sender email or a safe default."""
    sender = AppSetting.get("from_email") or current_app.config.get("MAIL_DEFAULT_SENDER")
    return sender or "noreply@example.com"


def _branding() -> dict[str, str | None]:
    """Return site title and optional absolute logo URL."""
    title = AppSetting.get("site_title", "VoteBuddy")
    logo_file = AppSetting.get("site_logo")
    logo_url = (
        url_for("static", filename=logo_file, _external=True) if logo_file else None
    )
    return {"site_title": title, "logo": logo_url}


def send_vote_invite(member: Member, token: str, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Send voting link to a member using Flask-Mail."""
    if member.email_opt_out:
        return
    link = url_for('voting.ballot_token', token=token, _external=True)
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Your voting link for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    objection_link = url_for('main.public_meeting_detail', meeting_id=meeting.id, _external=True)
    branding = _branding()
    notice_html = markdown_to_html(meeting.notice_md or "")
    msg.body = render_template(
        'email/invite.txt',
        member=member,
        meeting=meeting,
        link=link,
        notice_text=meeting.notice_md or "",
        objection_link=objection_link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        'email/invite.html',
        member=member,
        meeting=meeting,
        link=link,
        notice_html=notice_html,
        objection_link=objection_link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    try:
        ics = generate_stage_ics(meeting, 1)
    except Exception:
        ics = None
    if ics:
        msg.attach('stage1.ics', 'text/calendar', ics)
    mail.send(msg)
    _log_email(member, meeting, 'stage1_invite', test_mode)


def send_proxy_invite(proxy: Member, principal: Member, token: str, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Email proxy voting link to the proxy holder."""
    if proxy.email_opt_out:
        return
    link = url_for('voting.ballot_token', token=token, _external=True)
    unsubscribe = _unsubscribe_url(proxy)
    resubscribe = _resubscribe_url(proxy)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Proxy vote link for {meeting.title}",
        recipients=[proxy.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/proxy_invite.txt',
        proxy=proxy,
        principal=principal,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        'email/proxy_invite.html',
        proxy=proxy,
        principal=principal,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(proxy, meeting, 'proxy_invite', test_mode)


def send_stage2_invite(member: Member, token: str, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Email Stage 2 voting link to a member."""
    if member.email_opt_out:
        return
    link = url_for('voting.ballot_token', token=token, _external=True)
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    summary = carried_amendment_summary(meeting)
    if summary:
        results_link = None
    else:
        if meeting.early_public_results:
            results_link = url_for('main.public_stage1_results', meeting_id=meeting.id, _external=True)
        else:
            results_link = url_for('main.public_results', meeting_id=meeting.id, _external=True)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Stage 2 voting open for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/stage2_invite.txt',
        member=member,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        summary=summary,
        results_link=results_link,
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        'email/stage2_invite.html',
        member=member,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        summary=summary,
        results_link=results_link,
        test_mode=test_mode,
        **branding,
    )
    try:
        ics = generate_stage_ics(meeting, 2)
    except Exception:
        ics = None
    if ics:
        msg.attach('stage2.ics', 'text/calendar', ics)
    mail.send(msg)
    _log_email(member, meeting, 'stage2_invite', test_mode)

def send_runoff_invite(member: Member, token: str, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Email run-off voting link after Stage 1."""
    if member.email_opt_out:
        return
    link = url_for('voting.runoff_ballot', token=token, _external=True)
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Run-off vote for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/runoff_invite.txt',
        member=member,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        'email/runoff_invite.html',
        member=member,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    try:
        tmp_meeting = type('M', (), {
            'title': meeting.title,
            'opens_at_stage1': meeting.runoff_opens_at,
            'closes_at_stage1': meeting.runoff_closes_at,
        })()
        ics = generate_stage_ics(tmp_meeting, 1)
    except Exception:
        ics = None
    if ics:
        msg.attach('runoff.ics', 'text/calendar', ics)
    mail.send(msg)
    _log_email(member, meeting, 'runoff_invite', test_mode)


def send_stage1_reminder(member: Member, token: str, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Email reminder to cast Stage 1 vote."""
    if member.email_opt_out:
        return
    link = url_for('voting.ballot_token', token=token, _external=True)
    template_base = config_or_setting('REMINDER_TEMPLATE', 'email/reminder')
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Reminder: vote in {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    objection_link = url_for('main.public_meeting_detail', meeting_id=meeting.id, _external=True)
    branding = _branding()
    msg.body = render_template(
        f"{template_base}.txt",
        member=member,
        meeting=meeting,
        link=link,
        objection_link=objection_link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        f"{template_base}.html",
        member=member,
        meeting=meeting,
        link=link,
        objection_link=objection_link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'stage1_reminder', test_mode)


def send_stage2_reminder(member: Member, token: str, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Email reminder to cast Stage 2 vote."""
    if member.email_opt_out:
        return
    link = url_for('voting.ballot_token', token=token, _external=True)
    template_base = config_or_setting('STAGE2_REMINDER_TEMPLATE', 'email/stage2_reminder')
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    summary = carried_amendment_summary(meeting)
    if summary:
        results_link = None
    else:
        if meeting.early_public_results:
            results_link = url_for('main.public_stage1_results', meeting_id=meeting.id, _external=True)
        else:
            results_link = url_for('main.public_results', meeting_id=meeting.id, _external=True)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Reminder: vote in {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        f"{template_base}.txt",
        member=member,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        summary=summary,
        results_link=results_link,
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        f"{template_base}.html",
        member=member,
        meeting=meeting,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        summary=summary,
        results_link=results_link,
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'stage2_reminder', test_mode)


def send_vote_receipt(member: Member, meeting: Meeting, hashes: list[str], *, test_mode: bool = False) -> None:
    """Email a receipt containing vote hashes."""
    if member.email_opt_out:
        return
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Your vote receipt for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        "email/receipt.txt",
        member=member,
        meeting=meeting,
        hashes=hashes,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        "email/receipt.html",
        member=member,
        meeting=meeting,
        hashes=hashes,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'receipt', test_mode)

def send_quorum_failure(member: Member, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Notify a member that Stage 1 failed to reach quorum."""
    if member.email_opt_out:
        return
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Stage 1 vote void for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        "email/quorum_failure.txt",
        member=member,
        meeting=meeting,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        "email/quorum_failure.html",
        member=member,
        meeting=meeting,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'quorum_failure', test_mode)


def _shade_cell(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _styled_doc(title: str) -> Document:
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Gotham"
    normal.size = Pt(11)
    for level in ["Heading 2", "Heading 3"]:
        h = doc.styles[level].font
        h.name = "Gotham"
        h.color.rgb = RGBColor(0xDC, 0x07, 0x14)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    para = cell.paragraphs[0]
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade_cell(cell, "002D59")
    return doc


def _final_results_docx(meeting: Meeting) -> bytes:
    from ..models import Amendment, Motion, Vote
    amend_results = (
        Amendment.query.filter_by(meeting_id=meeting.id)
        .order_by(Amendment.order)
        .all()
    )
    motion_results = (
        Motion.query.filter_by(meeting_id=meeting.id)
        .order_by(Motion.ordering)
        .all()
    )

    doc = _styled_doc(f"{meeting.title} - Final Results")
    doc.add_heading("Carried Amendments", level=2)
    carried = []
    for amend in amend_results:
        counts = {
            "for": 0,
            "against": 0,
        }
        rows = (
            db.session.query(Vote.choice, func.count(Vote.id))
            .filter_by(amendment_id=amend.id)
            .filter(Vote.is_test.is_(False))
            .group_by(Vote.choice)
            .all()
        )
        for choice, count in rows:
            counts[choice] = count
        if counts.get("for", 0) > counts.get("against", 0):
            carried.append(amend)

    table_ca = doc.add_table(rows=1, cols=1)
    if carried:
        table_ca.rows[0].cells[0].text = carried[0].text_md or ""
        for idx, amend in enumerate(carried[1:], start=2):
            row = table_ca.add_row().cells
            row[0].text = amend.text_md or ""
            if idx % 2 == 0:
                _shade_cell(row[0], "F7F7F9")
    else:
        table_ca.rows[0].cells[0].text = "No amendments carried."

    doc.add_heading("Motion Outcomes", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Motion"
    hdr[1].text = "For"
    hdr[2].text = "Against"
    hdr[3].text = "Abstain"
    hdr[4].text = "Outcome"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for idx, motion in enumerate(motion_results, start=1):
        counts = {
            "for": 0,
            "against": 0,
            "abstain": 0,
        }
        rows = (
            db.session.query(Vote.choice, func.count(Vote.id))
            .filter_by(motion_id=motion.id)
            .filter(Vote.is_test.is_(False))
            .group_by(Vote.choice)
            .all()
        )
        for choice, count in rows:
            counts[choice] = count
        row = table.add_row().cells
        row[0].text = motion.title or "Motion"
        row[1].text = str(counts["for"])
        row[2].text = str(counts["against"])
        row[3].text = str(counts["abstain"])
        row[4].text = (motion.status or "?").capitalize()
        if idx % 2 == 0:
            for c in row:
                _shade_cell(c, "F7F7F9")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def send_final_results(member: Member, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Email certified results to a member."""
    if member.email_opt_out:
        return

    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    from ..utils import carried_amendment_summary, motion_results_summary

    ca_summary = carried_amendment_summary(meeting) or "No amendments carried."
    motion_summary = motion_results_summary(meeting)
    summary = f"{ca_summary}\n\nMotion Outcomes:\n{motion_summary}"

    results_link = (
        url_for('main.public_results', meeting_id=meeting.id, _external=True)
        if meeting.public_results
        else None
    )

    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Certified results for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    msg.body = render_template(
        'email/final_results.txt',
        member=member,
        meeting=meeting,
        summary=summary,
        results_link=results_link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
    )
    msg.html = render_template(
        'email/final_results.html',
        member=member,
        meeting=meeting,
        summary=summary,
        results_link=results_link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
    )

    if not results_link:
        doc_bytes = _final_results_docx(meeting)
        msg.attach(
            'final_results.docx',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            doc_bytes,
        )
    mail.send(msg)
    _log_email(member, meeting, 'final_results', test_mode)


def send_objection_confirmation(obj: AmendmentObjection, amendment: Amendment, meeting: Meeting, *, test_mode: bool = False) -> None:
    msg = Message(
        subject=("[TEST] " if test_mode else "") + "Confirm your objection",
        recipients=[obj.email],
        sender=_sender(),
    )
    link = url_for('meetings.confirm_objection', token=obj.token, _external=True)
    branding = _branding()
    msg.body = render_template(
        'email/objection_confirm.txt',
        link=link,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        'email/objection_confirm.html',
        link=link,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)


def send_board_notice(amendment: Amendment, meeting: Meeting, *, test_mode: bool = False) -> None:
    recipient = AppSetting.get("from_email", current_app.config.get("MAIL_DEFAULT_SENDER"))
    msg = Message(
        subject=("[TEST] " if test_mode else "") + "Objection threshold met",
        recipients=[recipient],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        "email/board_notice.txt",
        amendment=amendment,
        meeting=meeting,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        "email/board_notice.html",
        amendment=amendment,
        meeting=meeting,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)


def send_amendment_reinstated(amendment: Amendment, meeting: Meeting, *, test_mode: bool = False) -> None:
    coords = (
        User.query.join(Role, User.role_id == Role.id)
        .filter(Role.name == "meeting_coordinator", User.is_active.is_(True))
        .all()
    )
    recipients = [u.email for u in coords]
    if not recipients:
        return
    msg = Message(
        subject=("[TEST] " if test_mode else "") + "Amendment reinstated",
        recipients=recipients,
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        "email/amendment_reinstated.txt",
        amendment=amendment,
        meeting=meeting,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        "email/amendment_reinstated.html",
        amendment=amendment,
        meeting=meeting,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)

def send_motion_submission_alert(submission: MotionSubmission, meeting: Meeting, *, test_mode: bool = False) -> None:
    recipient = AppSetting.get('from_email', current_app.config.get('MAIL_DEFAULT_SENDER'))
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Motion submission for {meeting.title}",
        recipients=[recipient],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/motion_submission.txt',
        submission=submission,
        meeting=meeting,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        'email/motion_submission.html',
        submission=submission,
        meeting=meeting,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    try:
        mail.send(msg)
    except OSError as exc:
        current_app.logger.warning("Email send failed: %s", exc)


def send_amendment_submission_alert(submission: AmendmentSubmission, motion: Motion, meeting: Meeting, *, test_mode: bool = False) -> None:
    recipient = AppSetting.get('from_email', current_app.config.get('MAIL_DEFAULT_SENDER'))
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Amendment submission for {motion.title}",
        recipients=[recipient],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/amendment_submission.txt',
        submission=submission,
        motion=motion,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        'email/amendment_submission.html',
        submission=submission,
        motion=motion,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    try:
        mail.send(msg)
    except OSError as exc:
        current_app.logger.warning("Email send failed: %s", exc)


def notify_seconder_motion(seconder: Member, meeting: Meeting, *, test_mode: bool = False) -> None:
    if seconder.email_opt_out:
        return
    msg = Message(
        subject=("[TEST] " if test_mode else "") + "Motion seconded notification",
        recipients=[seconder.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        "email/seconder_notice.txt",
        meeting=meeting,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        "email/seconder_notice.html",
        meeting=meeting,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)


def notify_seconder_amendment(seconder: Member, meeting: Meeting, motion: Motion, *, test_mode: bool = False) -> None:
    if seconder.email_opt_out:
        return
    msg = Message(
        subject=("[TEST] " if test_mode else "") + "Amendment seconded notification",
        recipients=[seconder.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        "email/seconder_notice.txt",
        meeting=meeting,
        motion=motion,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        "email/seconder_notice.html",
        meeting=meeting,
        motion=motion,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)


def send_submission_invite(member: Member, meeting: Meeting, *, test_mode: bool = False) -> None:
    if member.email_opt_out:
        return
    token_obj, plain = SubmissionToken.create(
        member_id=member.id,
        meeting_id=meeting.id,
        salt=current_app.config["TOKEN_SALT"],
    )
    db.session.commit()
    link = url_for('submissions.submit_motion', token=plain, meeting_id=meeting.id, _external=True)
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Submit motions for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/submission_invite.txt',
        member=member,
        meeting=meeting,
        link=link,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        'email/submission_invite.html',
        member=member,
        meeting=meeting,
        link=link,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'submission_invite', test_mode)


def send_review_invite(member: Member, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Invite members to review motions and submit amendments."""
    if member.email_opt_out:
        return
    token_obj, plain = SubmissionToken.create(
        member_id=member.id,
        meeting_id=meeting.id,
        salt=current_app.config["TOKEN_SALT"],
    )
    db.session.commit()
    review_url = url_for('main.review_motions', token=plain, meeting_id=meeting.id, _external=True)
    link = url_for('submissions.submit_amendment_select', token=plain, meeting_id=meeting.id, _external=True)
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Review motions for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/review_invite.txt',
        member=member,
        meeting=meeting,
        review_url=review_url,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        'email/review_invite.html',
        member=member,
        meeting=meeting,
        review_url=review_url,
        link=link,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'review_invite', test_mode)


def send_amendment_review_invite(member: Member, meeting: Meeting, *, test_mode: bool = False) -> None:
    """Invite members to read submitted amendments and comment."""
    if member.email_opt_out:
        return
    token_obj, plain = SubmissionToken.create(
        member_id=member.id,
        meeting_id=meeting.id,
        salt=current_app.config["TOKEN_SALT"],
    )
    db.session.commit()
    review_url = url_for('main.review_motions', token=plain, meeting_id=meeting.id, _external=True)
    unsubscribe = _unsubscribe_url(member)
    resubscribe = _resubscribe_url(member)
    msg = Message(
        subject=("[TEST] " if test_mode else "") + f"Comment on amendments for {meeting.title}",
        recipients=[member.email],
        sender=_sender(),
    )
    branding = _branding()
    msg.body = render_template(
        'email/amendment_review_invite.txt',
        member=member,
        meeting=meeting,
        review_url=review_url,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    msg.html = render_template(
        'email/amendment_review_invite.html',
        member=member,
        meeting=meeting,
        review_url=review_url,
        unsubscribe_url=unsubscribe,
        resubscribe_url=resubscribe,
        why_text=config_or_setting('EMAIL_WHY_TEXT', DEFAULT_EMAIL_WHY_TEXT),
        test_mode=test_mode,
        **branding,
    )
    mail.send(msg)
    _log_email(member, meeting, 'amendment_review_invite', test_mode)

def send_password_reset(user: User, token: str, *, test_mode: bool = False) -> None:
    msg = Message(
        subject=("[TEST] " if test_mode else "") + "Reset your password",
        recipients=[user.email],
        sender=_sender(),
    )
    link = url_for("auth.reset_password", token=token, _external=True)
    branding = _branding()
    msg.body = render_template(
        "email/password_reset.txt",
        user=user,
        link=link,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    msg.html = render_template(
        "email/password_reset.html",
        user=user,
        link=link,
        why_text=config_or_setting("EMAIL_WHY_TEXT", DEFAULT_EMAIL_WHY_TEXT),
        **branding,
    )
    mail.send(msg)

