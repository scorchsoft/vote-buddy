from flask import (
    Blueprint,
    render_template,
    redirect,
    url_for,
    request,
    flash,
    abort,
    send_file,
    current_app,
    send_from_directory,
)
from flask_wtf import FlaskForm
from wtforms import TextAreaField, SubmitField
from wtforms.validators import DataRequired
from datetime import datetime, timedelta
from flask_login import login_required, current_user
from ..extensions import db
from ..models import (
    Meeting,
    Member,
    VoteToken,
    Amendment,
    AmendmentMerge,
    AmendmentConflict,
    AmendmentObjection,
    Motion,
    MotionOption,
    Vote,
    Runoff,
    AppSetting,
    MeetingFile,
    EmailSetting,
    EmailLog,    
    MotionVersion,
    AmendmentVersion,
    MotionSubmission,
    AmendmentSubmission,
    Comment,
)
from ..services.email import (
    send_vote_invite,
    send_stage2_invite,
    send_runoff_invite,
    send_quorum_failure,
    send_final_results,
    send_objection_confirmation,
    send_proxy_invite,
    send_submission_invite,
    send_review_invite,
    send_amendment_review_invite,
    send_stage1_reminder,
    auto_send_enabled,
    _branding,
)
from ..services import runoff
from ..services.audit import record_action
from ..comments import routes as comments
from ..permissions import permission_required
from .forms import (
    MeetingForm,
    MemberImportForm,
    AmendmentForm,
    MotionForm,
    ConflictForm,
    ObjectionForm,
    ManualEmailForm,
    ExtendStageForm,
    MotionChangeRequestForm,
    MeetingFileForm,
    Stage1TallyForm,
    Stage2TallyForm,
)
from ..voting.routes import (
    compile_motion_text,
    _amendment_form,
    _motion_form,
    _combined_form,
)
from ..utils import generate_stage_ics, markdown_to_html, carried_amendment_summary
import csv
import io
from uuid6 import uuid7
from sqlalchemy import func
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from ..utils import config_or_setting, append_motion_preferences
import os

bp = Blueprint("meetings", __name__, url_prefix="/meetings")


def _shade_cell(cell, color_hex: str) -> None:
    """Apply background colour to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _styled_doc(title: str, include_logo: bool) -> Document:
    """Return a Document with BP branding styles."""
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Gotham"
    normal.size = Pt(11)

    for level in ["Heading 2", "Heading 3"]:
        h = doc.styles[level].font
        h.name = "Gotham"
        h.color.rgb = RGBColor(0xDC, 0x07, 0x14)

    # header bar
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    para = cell.paragraphs[0]
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade_cell(cell, "002D59")

    if include_logo:
        logo_path = os.path.join(current_app.root_path, "..", "assets", "logo.png")
        footer = doc.sections[0].footer
        fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp_run = fp.add_run()
        try:
            fp_run.add_picture(logo_path, width=Inches(1))
        except Exception:
            pass

    return doc


def _calculate_default_times(agm_date: datetime) -> dict:
    """Calculate sensible default times for meeting phases based on AGM date."""
    cfg = current_app.config
    opens_at_stage2 = agm_date - timedelta(days=cfg.get("STAGE2_LENGTH_DAYS", 5))
    closes_at_stage1 = opens_at_stage2 - timedelta(days=cfg.get("STAGE_GAP_DAYS", 1))
    opens_at_stage1 = closes_at_stage1 - timedelta(
        days=cfg.get("STAGE1_LENGTH_DAYS", 7)
    )
    # Final notice comes after amendments close (3 days before Stage 1)
    notice_date = opens_at_stage1 - timedelta(days=cfg.get("NOTICE_PERIOD_DAYS", 3))
    motions_closes_at = notice_date - timedelta(
        days=cfg.get("MOTION_DEADLINE_GAP_DAYS", 0)  # Changed from 7 to 0 - motions open immediately
    )
    motions_opens_at = motions_closes_at - timedelta(
        days=cfg.get("MOTION_WINDOW_DAYS", 7)
    )
    # Initial notice comes before motions open (≥21 days before motions close)
    initial_notice_date = motions_closes_at - timedelta(days=21)
    # Amendments open same day as motions close (Day-One opening)
    amendments_opens_at = motions_closes_at
    # Amendments close after 5-day window
    amendments_closes_at = amendments_opens_at + timedelta(days=cfg.get("AMENDMENT_WINDOW_DAYS", 5))
    return {
        "initial_notice_date": initial_notice_date,
        "notice_date": notice_date,
        "opens_at_stage1": opens_at_stage1,
        "closes_at_stage1": closes_at_stage1,
        "opens_at_stage2": opens_at_stage2,
        "closes_at_stage2": agm_date,
        "motions_opens_at": motions_opens_at,
        "motions_closes_at": motions_closes_at,
        "amendments_opens_at": amendments_opens_at,
        "amendments_closes_at": amendments_closes_at,
    }


def _prefill_form_defaults(form: MeetingForm) -> None:
    """Set form defaults based on AGM date if fields are empty."""
    if form.closes_at_stage2.data:
        defaults = _calculate_default_times(form.closes_at_stage2.data)
        for field, value in defaults.items():
            if getattr(form, field).data is None:
                getattr(form, field).data = value


def _email_schedule(meeting: Meeting) -> dict[str, datetime | None]:
    """Return expected send times for key email types.

    Only include emails relevant to the meeting's ballot mode.
    """
    schedule = {
        "initial_notice": meeting.initial_notice_date,
        "submission_invite": meeting.motions_opens_at,
        "review_invite": meeting.amendments_opens_at,
        "amendment_review_invite": meeting.amendments_closes_at,
        "stage1_invite": meeting.notice_date,
        "stage1_reminder": (
            meeting.closes_at_stage1
            and meeting.closes_at_stage1
            - timedelta(
                hours=config_or_setting(
                    "REMINDER_HOURS_BEFORE_CLOSE", 6, parser=int
                )
            )
        ),
    }
    if meeting.ballot_mode == "two-stage":
        schedule.update(
            {
                "stage2_invite": meeting.opens_at_stage2,
                "stage2_reminder": (
                    meeting.closes_at_stage2
                    and meeting.closes_at_stage2
                    - timedelta(
                        hours=config_or_setting(
                            "STAGE2_REMINDER_HOURS_BEFORE_CLOSE", 6, parser=int
                        )
                    )
                ),
            }
        )
    return schedule


@bp.route("/")
@login_required
@permission_required("manage_meetings")
def list_meetings():
    q = request.args.get("q", "").strip()
    sort = request.args.get("sort", "title")
    direction = request.args.get("direction", "asc")

    page = request.args.get("page", 1, type=int)
    per_page = current_app.config.get("MEETINGS_PER_PAGE", 20)

    query = Meeting.query
    if q:
        search = f"%{q}%"
        query = query.filter(Meeting.title.ilike(search))

    if sort == "type":
        order_attr = Meeting.type
    elif sort == "status":
        order_attr = Meeting.status
    else:
        order_attr = Meeting.title

    query = query.order_by(
        order_attr.asc() if direction == "asc" else order_attr.desc()
    )

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    meetings = pagination.items

    # Only return partial template for specific HTMX requests (search/sort/pagination)
    # Not for general page navigation via hx-boost
    if (
        request.headers.get("HX-Request")
        and request.headers.get("HX-Target") == "meeting-table-body"
    ):
        template = "meetings/_meeting_rows.html"
    else:
        template = "meetings_list.html"

    return render_template(
        template,
        meetings=meetings,
        pagination=pagination,
        q=q,
        sort=sort,
        direction=direction,
    )


def _save_meeting(form: MeetingForm, meeting: Meeting | None = None) -> Meeting:
    """Populate Meeting from form and save."""
    if meeting is None:
        meeting = Meeting()
    _prefill_form_defaults(form)
    form.populate_obj(meeting)
    db.session.add(meeting)
    db.session.commit()
    if meeting.email_settings == []:
        for et in ["stage1_invite", "stage1_reminder", "stage2_invite", "stage2_reminder"]:
            db.session.add(EmailSetting(meeting_id=meeting.id, email_type=et, auto_send=True))
        db.session.commit()
    if (
        not meeting.opens_at_stage1
        and Amendment.query.filter_by(meeting_id=meeting.id).count() == 0
    ):
        meeting.status = "Pending Stage 2"
        db.session.commit()
    return meeting


@bp.route("/create", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def create_meeting():
    form = MeetingForm()
    notice_days = current_app.config.get("NOTICE_PERIOD_DAYS", 3)  # Updated from 14 to 3
    form.initial_notice_date.description = (
        "Basic meeting announcement sent early (≥21 days before motions close)."
    )
    form.notice_date.description = (
        f"Final notice with complete agenda; at least {notice_days} days before Stage 1 opens."
    )
    form.opens_at_stage1.description = f"At least {notice_days} days after final notice date."
    form.closes_at_stage1.description = "Must remain open for at least 5 days for e-ballots."  # Updated from 7 to 5
    form.opens_at_stage2.description = "At least 1 day after Stage 1 closes."
    form.closes_at_stage2.description = (
        "Final voting deadline; at least 5 days after Stage 2 opens."
    )
    if request.method == "GET":
        _prefill_form_defaults(form)
    if form.validate_on_submit():
        _save_meeting(form)
        return redirect(url_for("meetings.list_meetings"))
    return render_template("meetings/meetings_form.html", form=form)


@bp.route("/<int:meeting_id>/edit", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def edit_meeting(meeting_id):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    form = MeetingForm(obj=meeting)
    notice_days = current_app.config.get("NOTICE_PERIOD_DAYS", 3)  # Updated from 14 to 3
    form.initial_notice_date.description = (
        "Basic meeting announcement sent early (≥21 days before motions close)."
    )
    form.notice_date.description = (
        f"Final notice with complete agenda; at least {notice_days} days before Stage 1 opens."
    )
    form.opens_at_stage1.description = f"At least {notice_days} days after final notice date."
    form.closes_at_stage1.description = "Must remain open for at least 5 days for e-ballots."  # Updated from 7 to 5
    form.opens_at_stage2.description = "At least 1 day after Stage 1 closes."
    form.closes_at_stage2.description = (
        "Final voting deadline; at least 5 days after Stage 2 opens."
    )
    if request.method == "GET":
        _prefill_form_defaults(form)
    if form.validate_on_submit():
        _save_meeting(form, meeting)
        return redirect(url_for("meetings.edit_meeting", meeting_id=meeting_id))
    return render_template("meetings/meetings_form.html", form=form, meeting=meeting)


@bp.route("/<int:meeting_id>/clone")
@login_required
@permission_required("manage_meetings")
def clone_meeting(meeting_id: int):
    """Duplicate a meeting along with its motions and amendments."""
    src = db.session.get(Meeting, meeting_id)
    if src is None:
        abort(404)
    new_meeting = Meeting()
    for field in [
        "title",
        "type",
        "ballot_mode",
        "revoting_allowed",
        "status",
        "chair_notes_md",
        "quorum",
        "public_results",
        "early_public_results",
        "comments_enabled",
        "extension_reason",
        "results_doc_published",
        "results_doc_intro_md",
        "summary_md",
    ]:
        setattr(new_meeting, field, getattr(src, field))
    db.session.add(new_meeting)
    db.session.flush()

    motion_map: dict[int, int] = {}
    for motion in Motion.query.filter_by(meeting_id=src.id).all():
        new_motion = Motion(
            meeting_id=new_meeting.id,
            title=motion.title,
            text_md=motion.text_md,
            final_text_md=motion.final_text_md,
            category=motion.category,
            threshold=motion.threshold,
            ordering=motion.ordering,
            status=motion.status,
            withdrawn=motion.withdrawn,
        )
        db.session.add(new_motion)
        db.session.flush()
        for opt in motion.options:
            db.session.add(MotionOption(motion_id=new_motion.id, text=opt.text))
        motion_map[motion.id] = new_motion.id

    for amend in Amendment.query.filter_by(meeting_id=src.id).all():
        new_amend = Amendment(
            meeting_id=new_meeting.id,
            motion_id=motion_map.get(amend.motion_id),
            text_md=amend.text_md,
            order=amend.order,
            status=amend.status,
            reason=amend.reason,
            board_seconded=amend.board_seconded,
            tie_break_method=amend.tie_break_method,
        )
        db.session.add(new_amend)

    db.session.commit()
    return redirect(url_for("meetings.edit_meeting", meeting_id=new_meeting.id))


@bp.route("/<int:meeting_id>/import-members", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def import_members(meeting_id):
    """Upload a CSV of members and generate vote tokens."""

    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    form = MemberImportForm()
    if form.validate_on_submit():
        file_data = form.csv_file.data
        csv_text = file_data.read().decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(csv_text))
        expected = ["member_id", "name", "email", "proxy_for"]
        if reader.fieldnames != expected:
            flash("CSV headers must be: " + ", ".join(expected), "error")
            return render_template(
                "meetings/import_members.html", form=form, meeting=meeting
            )

        seen_emails: set[str] = set()
        seen_numbers: set[str] = set()
        tokens_to_send: list[tuple[Member, str]] = []
        proxy_tokens: list[tuple[Member, Member, str]] = []
        for idx, row in enumerate(reader, start=2):
            name = row["name"].strip()
            email = row["email"].strip().lower()
            number = (row.get("member_id") or "").strip()

            if not name:
                flash(f"Row {idx}: name is required", "error")
                return render_template(
                    "meetings/import_members.html", form=form, meeting=meeting
                )
            if not email or "@" not in email:
                flash(f"Row {idx}: invalid email: {email}", "error")
                return render_template(
                    "meetings/import_members.html", form=form, meeting=meeting
                )
            if (
                email in seen_emails
                or Member.query.filter_by(meeting_id=meeting.id, email=email).first()
            ):
                flash(f"Duplicate email: {email}", "error")
                return render_template(
                    "meetings/import_members.html", form=form, meeting=meeting
                )
            if number and (
                number in seen_numbers
                or Member.query.filter_by(
                    meeting_id=meeting.id, member_number=number
                ).first()
            ):
                flash(f"Duplicate member ID: {number}", "error")
                return render_template(
                    "meetings/import_members.html", form=form, meeting=meeting
                )
            seen_emails.add(email)
            if number:
                seen_numbers.add(number)

            member = Member(
                meeting_id=meeting.id,
                member_number=number,
                name=name,
                email=email,
                proxy_for=(row.get("proxy_for") or "").strip() or None,
            )
            db.session.add(member)
            db.session.flush()
            if meeting.ballot_mode != "in-person":
                token_obj, plain = VoteToken.create(
                    member_id=member.id,
                    stage=1,
                    salt=current_app.config["TOKEN_SALT"],
                )
                tokens_to_send.append((member, plain))

        db.session.commit()
        # create proxy tokens after all members exist
        members = Member.query.filter_by(meeting_id=meeting.id).all()
        for proxy in members:
            if proxy.proxy_for:
                try:
                    target = db.session.get(Member, int(proxy.proxy_for))
                except (ValueError, TypeError):
                    target = None
                if target and meeting.ballot_mode != "in-person":
                    token_obj, plain = VoteToken.create(
                        member_id=target.id,
                        stage=1,
                        salt=current_app.config["TOKEN_SALT"],
                        proxy_holder_id=proxy.id,
                    )
                    proxy_tokens.append((proxy, target, plain))
        db.session.commit()

        if auto_send_enabled(meeting, 'stage1_invite'):
            for m, t in tokens_to_send:
                send_vote_invite(m, t, meeting)
            for p, target, tok in proxy_tokens:
                send_proxy_invite(p, target, tok, meeting)
        else:
            flash("Automatic emails disabled - use manual send", "warning")
        flash("Members imported successfully", "success")
        return redirect(url_for("meetings.list_meetings"))

    return render_template("meetings/import_members.html", form=form, meeting=meeting)


@bp.route("/sample-members.csv")
@login_required
@permission_required("manage_meetings")
def download_sample_csv():
    """Serve a template CSV for member uploads."""
    path = os.path.join(current_app.root_path, "static", "sample_members.csv")
    return send_file(
        path,
        mimetype="text/csv",
        as_attachment=True,
        download_name="sample_members.csv",
    )


@bp.route("/<int:meeting_id>/members")
@login_required
@permission_required("manage_meetings")
def list_members(meeting_id: int):
    """View members uploaded to a meeting."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)

    q = request.args.get("q", "").strip()
    voted = request.args.get("voted", "any")
    sort = request.args.get("sort", "name")
    direction = request.args.get("direction", "asc")

    query = Member.query.filter_by(meeting_id=meeting.id)
    if q:
        search = f"%{q}%"
        query = query.filter(
            db.or_(Member.name.ilike(search), Member.email.ilike(search))
        )

    if voted in {"yes", "no"}:
        voted_subq = db.session.query(VoteToken.id).filter(
            VoteToken.member_id == Member.id,
            VoteToken.used_at.isnot(None),
        )
        if voted == "yes":
            query = query.filter(voted_subq.exists())
        else:
            query = query.filter(~voted_subq.exists())

    order_attr = Member.email if sort == "email" else Member.name
    query = query.order_by(
        order_attr.asc() if direction == "asc" else order_attr.desc()
    )

    page = request.args.get("page", 1, type=int)
    per_page = current_app.config.get("MEMBERS_PER_PAGE", 20)
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    members = []
    for m in pagination.items:
        voted_flag = (
            VoteToken.query.filter_by(member_id=m.id)
            .filter(VoteToken.used_at.isnot(None))
            .count()
            > 0
        )
        m.voted = voted_flag
        members.append(m)

    now = datetime.utcnow()
    runoff_exists = Runoff.query.filter_by(meeting_id=meeting.id).count() > 0
    email_opts = {
        "stage1_invite": meeting.opens_at_stage1
        and now >= meeting.opens_at_stage1
        and (meeting.opens_at_stage2 is None or now < meeting.opens_at_stage2),
        "stage1_reminder": meeting.opens_at_stage1
        and now >= meeting.opens_at_stage1
        and (meeting.closes_at_stage1 is None or now <= meeting.closes_at_stage1),
        "runoff_invite": runoff_exists
        and meeting.stage1_closed_at
        and meeting.opens_at_stage2
        and meeting.stage1_closed_at <= now < meeting.opens_at_stage2,
        "stage2_invite": meeting.ballot_mode == "two-stage"
        and meeting.opens_at_stage2
        and now >= meeting.opens_at_stage2
        and (meeting.closes_at_stage2 is None or now <= meeting.closes_at_stage2),
        "stage2_reminder": meeting.ballot_mode == "two-stage"
        and meeting.opens_at_stage2
        and now >= meeting.opens_at_stage2
        and (meeting.closes_at_stage2 is None or now <= meeting.closes_at_stage2),
        "submission_invite": meeting.motions_opens_at
        and now >= meeting.motions_opens_at
        and (meeting.motions_closes_at is None or now <= meeting.motions_closes_at),
        "review_invite": meeting.amendments_opens_at
        and now >= meeting.amendments_opens_at
        and (meeting.amendments_closes_at is None or now <= meeting.amendments_closes_at),
        "amendment_review_invite": meeting.amendments_closes_at
        and now >= meeting.amendments_closes_at
        and (meeting.opens_at_stage1 is None or now < meeting.opens_at_stage1),
        "final_results": meeting.status == "Completed",
    }

    if (
        request.headers.get("HX-Request")
        and request.headers.get("HX-Target") == "member-table-body"
    ):
        template = "meetings/_member_rows.html"
    else:
        template = "meetings/members.html"

    return render_template(
        template,
        meeting=meeting,
        members=members,
        pagination=pagination,
        q=q,
        voted=voted,
        sort=sort,
        direction=direction,
        now=now,
        runoff_exists=runoff_exists,
        email_opts=email_opts,
    )


@bp.route("/<int:meeting_id>/members/<int:member_id>/delete", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def delete_member(meeting_id: int, member_id: int):
    """Remove a single member from a meeting."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    member = Member.query.filter_by(id=member_id, meeting_id=meeting.id).first_or_404()
    Vote.query.filter_by(member_id=member.id).delete()
    VoteToken.query.filter_by(member_id=member.id).delete()
    db.session.delete(member)
    db.session.commit()
    flash("Member removed", "success")
    return redirect(url_for("meetings.list_members", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/members/delete-all", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def delete_all_members(meeting_id: int):
    """Remove all members from a meeting."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    member_ids = [m.id for m in Member.query.filter_by(meeting_id=meeting.id).all()]
    if member_ids:
        Vote.query.filter(Vote.member_id.in_(member_ids)).delete(
            synchronize_session=False
        )
        VoteToken.query.filter(VoteToken.member_id.in_(member_ids)).delete(
            synchronize_session=False
        )
        Member.query.filter(Member.id.in_(member_ids)).delete(synchronize_session=False)
        db.session.commit()
    flash("All members removed", "success")
    return redirect(url_for("meetings.list_members", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/members.csv")
@login_required
@permission_required("manage_meetings")
def members_csv(meeting_id: int):
    """Download a CSV list of members and their amendment votes."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)

    amendments = (
        Amendment.query.filter_by(meeting_id=meeting.id).order_by(Amendment.order).all()
    )

    output = io.StringIO()
    headers = ["member_id", "name", "email", "voted"] + [
        f"amend_{a.id}" for a in amendments
    ]
    writer = csv.writer(output)
    writer.writerow(headers)

    for m in Member.query.filter_by(meeting_id=meeting.id).order_by(Member.name).all():
        voted = (
            VoteToken.query.filter_by(member_id=m.id)
            .filter(VoteToken.used_at.isnot(None))
            .count()
            > 0
        )
        row = [m.member_number or "", m.name, m.email, "yes" if voted else "no"]
        for a in amendments:
            v = Vote.query.filter_by(member_id=m.id, amendment_id=a.id).first()
            row.append(v.choice if v else "")
        writer.writerow(row)

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode()),
        mimetype="text/csv",
        as_attachment=True,
        download_name="members.csv",
    )


@bp.route("/<int:meeting_id>/files", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def meeting_files(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    form = MeetingFileForm()
    if form.validate_on_submit():
        data = form.file.data
        root = current_app.config.get(
            "UPLOAD_FOLDER", os.path.join(current_app.instance_path, "files")
        )
        meeting_dir = os.path.join(root, str(meeting.id))
        os.makedirs(meeting_dir, exist_ok=True)
        filename = f"{uuid7()}.pdf"
        path = os.path.join(meeting_dir, filename)
        data.save(path)
        mf = MeetingFile(
            meeting_id=meeting.id,
            filename=filename,
            title=form.title.data,
            description=form.description.data,
        )
        db.session.add(mf)
        db.session.commit()
        flash("File uploaded", "success")
        return redirect(url_for("meetings.meeting_files", meeting_id=meeting.id))
    files = MeetingFile.query.filter_by(meeting_id=meeting.id).all()
    return render_template(
        "meetings/meeting_files.html", meeting=meeting, form=form, files=files
    )


@bp.route("/<int:meeting_id>/meeting-overview")
@login_required
@permission_required("manage_meetings")
def meeting_overview(meeting_id):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    motions = (
        Motion.query.filter_by(meeting_id=meeting.id)
        .order_by(Motion.ordering)
        .all()
    )
    amendments_count = Amendment.query.filter_by(meeting_id=meeting.id).count()
    counts = (
        db.session.query(Amendment.motion_id, func.count(Amendment.id))
        .filter(Amendment.meeting_id == meeting.id)
        .group_by(Amendment.motion_id)
        .all()
    )
    amendment_counts = {mid: c for mid, c in counts}
    votes_cast = meeting.stage1_votes_count()
    from datetime import datetime

    steps = [
        ("Initial Notice", meeting.initial_notice_date),
        ("Motions Open", meeting.motions_opens_at),
        ("Motions Close", meeting.motions_closes_at),
        ("Amendments Open", meeting.amendments_opens_at),
        ("Amendments Close", meeting.amendments_closes_at),
        ("Final Notice", meeting.notice_date),
        ("Stage 1 Opens", meeting.opens_at_stage1),
        ("Stage 1 Closes", meeting.closes_at_stage1),
        ("Stage 2 Opens", meeting.opens_at_stage2),
        ("AGM Date", meeting.closes_at_stage2),
    ]
    dates = [d for _, d in steps if d]
    timeline_start = min(dates) if dates else None
    timeline_end = max(dates) if dates else None
    schedule = _email_schedule(meeting)
    settings = {s.email_type: s for s in meeting.email_settings}
    members_count = Member.query.filter_by(meeting_id=meeting.id).count()
    pending_motions = MotionSubmission.query.filter_by(meeting_id=meeting.id).count()
    pending_amendments = (
        AmendmentSubmission.query.join(Motion, AmendmentSubmission.motion_id == Motion.id)
        .filter(Motion.meeting_id == meeting.id)
        .count()
    )
    files_count = MeetingFile.query.filter_by(meeting_id=meeting.id).count()

    return render_template(
        "meetings/meeting_overview.html",
        meeting=meeting,
        motions=motions,
        amendments_count=amendments_count,
        amendment_counts=amendment_counts,
        votes_cast=votes_cast,
        timeline_steps=steps,
        timeline_start=timeline_start,
        timeline_end=timeline_end,
        schedule=schedule,
        settings=settings,
        now=datetime.utcnow(),
        members_count=members_count,
        pending_motions=pending_motions,
        pending_amendments=pending_amendments,
        files_count=files_count,
    )


@bp.route("/<int:meeting_id>/motions")
@login_required
@permission_required("manage_meetings")
def list_motions(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    motions = (
        Motion.query.filter_by(meeting_id=meeting.id)
        .order_by(Motion.ordering)
        .all()
    )
    amendments = (
        Amendment.query.filter_by(meeting_id=meeting.id)
        .order_by(Amendment.order)
        .all()
    )
    amendments_by_motion: dict[int, list[Amendment]] = {}
    for amend in amendments:
        amendments_by_motion.setdefault(amend.motion_id, []).append(amend)
    return render_template(
        "meetings/motions_list.html",
        meeting=meeting,
        motions=motions,
        amendments_by_motion=amendments_by_motion,
    )


@bp.route("/<int:meeting_id>/motions/batch-edit", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def batch_edit_motions(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    motions = Motion.query.filter_by(meeting_id=meeting.id).order_by(Motion.ordering).all()
    amendments = Amendment.query.filter_by(meeting_id=meeting.id).order_by(Amendment.order).all()
    if request.method == "POST":
        for motion in motions:
            prefix = f"motion-{motion.id}-"
            if request.form.get(prefix + "title") is not None:
                db.session.add(
                    MotionVersion(
                        motion_id=motion.id,
                        title=motion.title,
                        text_md=motion.text_md,
                        final_text_md=motion.final_text_md,
                        proposer_id=motion.proposer_id,
                        seconder_id=motion.seconder_id,
                        board_proposed=motion.board_proposed,
                        board_seconded=motion.board_seconded,
                    )
                )
                motion.title = request.form.get(prefix + "title")
                motion.text_md = request.form.get(prefix + "text_md")
                motion.final_text_md = request.form.get(prefix + "final_text_md")
                motion.proposer_id = request.form.get(prefix + "proposer_id") or None
                motion.seconder_id = request.form.get(prefix + "seconder_id") or None
                motion.board_proposed = bool(request.form.get(prefix + "board_proposed"))
                motion.board_seconded = bool(request.form.get(prefix + "board_seconded"))
        for amend in amendments:
            prefix = f"amend-{amend.id}-"
            if request.form.get(prefix + "text_md") is not None:
                db.session.add(
                    AmendmentVersion(
                        amendment_id=amend.id,
                        text_md=amend.text_md,
                        proposer_id=amend.proposer_id,
                        seconder_id=amend.seconder_id,
                        board_proposed=amend.board_proposed,
                        board_seconded=amend.board_seconded,
                    )
                )
                amend.text_md = request.form.get(prefix + "text_md")
                amend.proposer_id = request.form.get(prefix + "proposer_id") or None
                amend.seconder_id = request.form.get(prefix + "seconder_id") or None
                amend.board_proposed = bool(request.form.get(prefix + "board_proposed"))
                amend.board_seconded = bool(request.form.get(prefix + "board_seconded"))
        if request.form.get("new_motion_title"):
            ordering = len(motions) + 1
            m = Motion(
                meeting_id=meeting.id,
                title=request.form.get("new_motion_title"),
                text_md=request.form.get("new_motion_text_md"),
                final_text_md=request.form.get("new_motion_final_text_md") or None,
                category=request.form.get("new_motion_category", "motion"),
                threshold=request.form.get("new_motion_threshold", "normal"),
                ordering=ordering,
                proposer_id=request.form.get("new_motion_proposer_id") or None,
                seconder_id=request.form.get("new_motion_seconder_id") or None,
                board_proposed=bool(request.form.get("new_motion_board_proposed")),
                board_seconded=bool(request.form.get("new_motion_board_seconded")),
            )
            db.session.add(m)
        if request.form.get("new_amend_text_md"):
            motion_id_raw = request.form.get("new_amend_motion_id", "").strip()
            if not motion_id_raw.isdigit():
                flash("Select the motion this amendment applies to.", "error")
                return redirect(
                    url_for("meetings.batch_edit_motions", meeting_id=meeting.id)
                )
            m_id = int(motion_id_raw)
            order = Amendment.query.filter_by(motion_id=m_id).count() + 1
            a = Amendment(
                meeting_id=meeting.id,
                motion_id=m_id,
                text_md=request.form.get("new_amend_text_md"),
                order=order,
                proposer_id=request.form.get("new_amend_proposer_id") or None,
                seconder_id=request.form.get("new_amend_seconder_id") or None,
                board_proposed=bool(request.form.get("new_amend_board_proposed")),
                board_seconded=bool(request.form.get("new_amend_board_seconded")),
            )
            db.session.add(a)
        db.session.commit()
        record_action("batch_edit_motions", f"meeting_id={meeting.id}")
        flash("Changes saved", "success")
        return redirect(url_for("meetings.batch_edit_motions", meeting_id=meeting.id))
    return render_template(
        "meetings/batch_edit_motions.html",
        meeting=meeting,
        motions=motions,
        amendments=amendments,
    )


@bp.route("/<int:meeting_id>/motions/create", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def create_motion(meeting_id):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    form = MotionForm()
    clerical_text = config_or_setting("CLERICAL_TEXT", "")
    move_text = config_or_setting("MOVE_TEXT", "")
    if form.validate_on_submit():
        ordering = Motion.query.filter_by(meeting_id=meeting.id).count() + 1
        text_md = append_motion_preferences(
            form.text_md.data,
            form.allow_clerical.data,
            form.allow_move.data,
            clerical_text,
            move_text,
        )
        motion = Motion(
            meeting_id=meeting.id,
            title=form.title.data,
            text_md=text_md,
            category=form.category.data,
            threshold=form.threshold.data,
            ordering=ordering,
        )
        db.session.add(motion)
        db.session.flush()
        if form.category.data == "multiple_choice" and form.options.data:
            for line in form.options.data.splitlines():
                text = line.strip()
                if text:
                    db.session.add(MotionOption(motion_id=motion.id, text=text))
        db.session.commit()
        return redirect(url_for("meetings.list_motions", meeting_id=meeting.id))
    return render_template(
        "meetings/motion_form.html",
        form=form,
        motion=None,
        clerical_text=clerical_text,
        move_text=move_text,
    )


@bp.route("/motions/<int:motion_id>")
def view_motion(motion_id):
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    if not motion.is_published:
        if not current_user.is_authenticated or not current_user.has_permission(
            "manage_meetings"
        ):
            abort(404)
    amendments = (
        Amendment.query.filter_by(motion_id=motion.id).order_by(Amendment.order).all()
    )
    conflicts = (
        AmendmentConflict.query.join(
            Amendment, Amendment.id == AmendmentConflict.amendment_a_id
        )
        .filter(Amendment.motion_id == motion.id)
        .all()
    )
    token = request.args.get("token")
    if token:
        token_obj = SubmissionToken.verify(token, current_app.config["TOKEN_SALT"])
        if not token_obj or token_obj.meeting_id != motion.meeting_id:
            token = None
    return render_template(
        "meetings/view_motion.html",
        motion=motion,
        amendments=amendments,
        conflicts=conflicts,
        token=token,
    )


@bp.route("/motions/<int:motion_id>/request-change", methods=["GET", "POST"])
@login_required
def request_motion_change(motion_id: int):
    """Allow a user to request withdrawal or major edit of a motion."""
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    meeting = db.session.get(Meeting, motion.meeting_id)
    if meeting is None:
        abort(404)
    if (
        meeting.opens_at_stage1
        and datetime.utcnow() > meeting.opens_at_stage1 - timedelta(days=7)
    ):
        flash(
            "Change requests must be made at least 7 days before Stage 1 opens.",
            "error",
        )
        return redirect(url_for("meetings.view_motion", motion_id=motion.id))
    form = MotionChangeRequestForm()
    if request.method == "GET":
        form.text_md.data = motion.text_md
    if form.validate_on_submit():
        if form.text_md.data and form.text_md.data != motion.text_md:
            motion.text_md = form.text_md.data
            motion.modified_at = datetime.utcnow()
            motion.withdrawal_requested_at = datetime.utcnow()
        if form.withdraw.data:
            motion.withdrawal_requested_at = datetime.utcnow()
        db.session.commit()
        flash("Request submitted", "success")
        return redirect(url_for("meetings.view_motion", motion_id=motion.id))
    return render_template("meetings/motion_change_form.html", form=form, motion=motion)


@bp.route("/motions/<int:motion_id>/approve-change/<actor>", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def approve_motion_change(motion_id: int, actor: str):
    """Approve a withdrawal/edit request as chair or board."""
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    now = datetime.utcnow()
    if actor == "chair":
        motion.chair_approved_at = now
    elif actor == "board":
        motion.board_approved_at = now
    else:
        abort(404)
    if motion.chair_approved_at and motion.board_approved_at:
        if motion.withdrawal_requested_at:
            motion.withdrawn = True
    db.session.commit()
    flash("Request approved", "success")
    return redirect(url_for("meetings.view_motion", motion_id=motion.id))


@bp.route("/motions/<int:motion_id>/reject-change", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def reject_motion_change(motion_id: int):
    """Reject a withdrawal/edit request."""
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    motion.withdrawal_requested_at = None
    motion.chair_approved_at = None
    motion.board_approved_at = None
    db.session.commit()
    flash("Request rejected", "success")
    return redirect(url_for("meetings.view_motion", motion_id=motion.id))


@bp.route("/motions/<int:motion_id>/amendments/add", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def add_amendment(motion_id):
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    form = AmendmentForm()
    members = (
        Member.query.filter_by(meeting_id=motion.meeting_id).order_by(Member.name).all()
    )
    choices = [(m.id, m.name) for m in members]
    form.proposer_id.choices = choices
    form.seconder_id.choices = [(0, "")] + choices
    if form.validate_on_submit():
        meeting = db.session.get(Meeting, motion.meeting_id)
        deadline = None
        if meeting.opens_at_stage1:
            deadline = meeting.opens_at_stage1 - timedelta(days=21)
        elif meeting.notice_date:
            notice_days = current_app.config.get("NOTICE_PERIOD_DAYS", 3)  # Updated from 14 to 3
            deadline = meeting.notice_date + timedelta(days=notice_days)

        if deadline and datetime.utcnow() > deadline:
            flash("Amendment deadline has passed.", "error")
            return render_template(
                "meetings/amendment_form.html", form=form, motion=motion
            )

        if not form.seconder_id.data and not form.board_seconded.data:
            flash("Select a seconder or confirm board approval.", "error")
            return render_template(
                "meetings/amendment_form.html", form=form, motion=motion
            )
        if form.seconder_id.data and form.proposer_id.data == form.seconder_id.data:
            flash("Proposer cannot second their own amendment.", "error")
            return render_template(
                "meetings/amendment_form.html", form=form, motion=motion
            )

        count = Amendment.query.filter_by(
            motion_id=motion.id,
            proposer_id=form.proposer_id.data,
        ).count()
        if count >= 3:
            flash("A member may propose at most three amendments per motion.", "error")
            return render_template(
                "meetings/amendment_form.html", form=form, motion=motion
            )

        order = Amendment.query.filter_by(motion_id=motion.id).count() + 1
        amendment = Amendment(
            meeting_id=motion.meeting_id,
            motion_id=motion.id,
            text_md=form.text_md.data,
            order=order,
            proposer_id=form.proposer_id.data,
            seconder_id=(
                (form.seconder_id.data or None)
                if not form.board_seconded.data
                else None
            ),
            board_seconded=form.board_seconded.data,
            seconded_method=form.seconded_method.data or None,
            seconded_at=datetime.utcnow() if form.seconded_method.data else None,
        )
        db.session.add(amendment)
        db.session.flush()

        combine_param = request.args.get("combine")
        if combine_param:
            ids = [int(i) for i in combine_param.split(",") if i.isdigit()]
            for src_id in ids:
                db.session.add(
                    AmendmentMerge(combined_id=amendment.id, source_id=src_id)
                )
                src = db.session.get(Amendment, src_id)
                if src:
                    src.status = "merged"

        db.session.commit()
        return redirect(url_for("meetings.view_motion", motion_id=motion.id))
    return render_template("meetings/amendment_form.html", form=form, motion=motion)


@bp.route("/amendments/<int:amendment_id>/edit", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def edit_amendment(amendment_id: int):
    """Edit an existing amendment."""
    amendment = db.session.get(Amendment, amendment_id)
    if amendment is None:
        abort(404)
    motion = db.session.get(Motion, amendment.motion_id)
    if motion is None:
        abort(404)
    meeting = db.session.get(Meeting, amendment.meeting_id)

    form = AmendmentForm(obj=amendment)
    members = Member.query.filter_by(meeting_id=meeting.id).order_by(Member.name).all()
    choices = [(m.id, m.name) for m in members]
    form.proposer_id.choices = choices
    form.seconder_id.choices = [(0, "")] + choices

    if form.validate_on_submit():
        if meeting.opens_at_stage1:
            deadline = meeting.opens_at_stage1 - timedelta(days=21)
            if datetime.utcnow() > deadline:
                flash("Amendment deadline has passed.", "error")
                return render_template(
                    "meetings/amendment_form.html",
                    form=form,
                    motion=motion,
                    amendment=amendment,
                )

        if not form.seconder_id.data and not form.board_seconded.data:
            flash("Select a seconder or confirm board approval.", "error")
            return render_template(
                "meetings/amendment_form.html",
                form=form,
                motion=motion,
                amendment=amendment,
            )
        if form.seconder_id.data and form.proposer_id.data == form.seconder_id.data:
            flash("Proposer cannot second their own amendment.", "error")
            return render_template(
                "meetings/amendment_form.html",
                form=form,
                motion=motion,
                amendment=amendment,
            )

        amendment.text_md = form.text_md.data
        amendment.proposer_id = form.proposer_id.data
        amendment.seconder_id = (
            form.seconder_id.data or None if not form.board_seconded.data else None
        )
        amendment.board_seconded = form.board_seconded.data
        amendment.seconded_method = form.seconded_method.data or None
        amendment.seconded_at = datetime.utcnow() if form.seconded_method.data else None
        db.session.commit()
        flash("Amendment updated", "success")
        return redirect(url_for("meetings.view_motion", motion_id=motion.id))

    return render_template(
        "meetings/amendment_form.html",
        form=form,
        motion=motion,
        amendment=amendment,
    )


@bp.route("/amendments/<int:amendment_id>/delete", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def delete_amendment(amendment_id: int):
    """Delete an amendment."""
    amendment = db.session.get(Amendment, amendment_id)
    if amendment is None:
        abort(404)
    meeting = db.session.get(Meeting, amendment.meeting_id)
    motion_id = amendment.motion_id

    if meeting.opens_at_stage1:
        deadline = meeting.opens_at_stage1 - timedelta(days=21)
        if datetime.utcnow() > deadline:
            flash("Amendment deadline has passed.", "error")
            return redirect(url_for("meetings.view_motion", motion_id=motion_id))

    db.session.delete(amendment)
    db.session.commit()
    flash("Amendment deleted", "success")
    return redirect(url_for("meetings.view_motion", motion_id=motion_id))


@bp.route("/amendments/<int:amendment_id>/reject", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def reject_amendment(amendment_id: int):
    """Mark an amendment as rejected."""
    amendment = db.session.get(Amendment, amendment_id)
    if amendment is None:
        abort(404)
    amendment.status = "rejected"
    amendment.reason = request.form.get("reason")
    db.session.commit()
    flash("Amendment marked as rejected", "success")
    return redirect(url_for("meetings.view_motion", motion_id=amendment.motion_id))


@bp.route("/amendments/<int:amendment_id>/mark-merged", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def mark_amendment_merged(amendment_id: int):
    """Mark an amendment as merged into another."""
    amendment = db.session.get(Amendment, amendment_id)
    if amendment is None:
        abort(404)
    amendment.status = "merged"
    amendment.reason = request.form.get("reason")
    db.session.commit()
    flash("Amendment marked as merged", "success")
    return redirect(url_for("meetings.view_motion", motion_id=amendment.motion_id))


@bp.route("/motions/<int:motion_id>/toggle-publish", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def toggle_motion_publish(motion_id: int):
    """Publish or unpublish a motion."""
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    motion.is_published = not motion.is_published
    db.session.commit()
    record_action("toggle_motion_publish", f"motion_id={motion.id}")
    return redirect(request.referrer or url_for("meetings.view_motion", motion_id=motion.id))


@bp.route("/<int:meeting_id>/member-search")
def member_search(meeting_id: int):
    """Return member options filtered by query."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    q = request.args.get("q", "").strip()
    query = Member.query.filter_by(meeting_id=meeting.id)
    if q:
        search = f"%{q}%"
        query = query.filter(
            db.or_(
                Member.name.ilike(search),
                Member.email.ilike(search),
                Member.member_number.ilike(search),
            )
        )
    members = query.order_by(Member.name).limit(20).all()
    return render_template("meetings/_member_options.html", members=members)


@bp.route("/amendments/<int:amendment_id>/object", methods=["GET", "POST"])
def submit_objection(amendment_id: int):
    """Allow a member to submit an objection."""
    amendment = db.session.get(Amendment, amendment_id)
    if amendment is None:
        abort(404)
    meeting = db.session.get(Meeting, amendment.meeting_id)
    if meeting is None:
        abort(404)
    form = ObjectionForm()
    if form.validate_on_submit():
        try:
            member_id = int(form.member_id.data)
        except (TypeError, ValueError):
            flash("Invalid member", "error")
            return render_template(
                "meetings/objection_form.html", form=form, amendment=amendment
            )

        member = Member.query.filter_by(id=member_id, meeting_id=meeting.id).first()
        if not member:
            flash("Member not found", "error")
            return render_template(
                "meetings/objection_form.html", form=form, amendment=amendment
            )

        obj = AmendmentObjection(
            amendment_id=amendment.id,
            member_id=member.id,
            email=form.email.data.strip().lower(),
            token=str(uuid7()),
        )
        db.session.add(obj)
        db.session.commit()
        send_objection_confirmation(obj, amendment, meeting)
        flash("Check your email to confirm the objection", "success")
        return redirect(url_for("meetings.view_motion", motion_id=amendment.motion_id))
    return render_template(
        "meetings/objection_form.html", form=form, amendment=amendment
    )


@bp.get("/objection/confirm/<token>")
def confirm_objection(token: str):
    obj = AmendmentObjection.query.filter_by(token=token).first_or_404()
    if not obj.confirmed_at:
        obj.confirmed_at = datetime.utcnow()
        obj.deadline_first = obj.created_at + timedelta(days=5)
        db.session.commit()
    return render_template("meetings/objection_confirmed.html", objection=obj)


@bp.route("/motions/<int:motion_id>/conflicts", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def manage_conflicts(motion_id: int):
    motion = db.session.get(Motion, motion_id)
    if motion is None:
        abort(404)
    amendments = (
        Amendment.query.filter_by(motion_id=motion.id).order_by(Amendment.order).all()
    )
    form = ConflictForm()
    choices = [(a.id, f"A{a.order}") for a in amendments]
    form.amendment_a_id.choices = choices
    form.amendment_b_id.choices = choices

    if form.validate_on_submit():
        a_id = form.amendment_a_id.data
        b_id = form.amendment_b_id.data
        if a_id != b_id:
            exists = (
                AmendmentConflict.query.filter_by(meeting_id=motion.meeting_id)
                .filter(
                    (
                        (AmendmentConflict.amendment_a_id == a_id)
                        & (AmendmentConflict.amendment_b_id == b_id)
                    )
                    | (
                        (AmendmentConflict.amendment_a_id == b_id)
                        & (AmendmentConflict.amendment_b_id == a_id)
                    )
                )
                .first()
            )
            if not exists:
                db.session.add(
                    AmendmentConflict(
                        meeting_id=motion.meeting_id,
                        amendment_a_id=a_id,
                        amendment_b_id=b_id,
                    )
                )
                db.session.commit()
                flash("Conflict recorded", "success")
            return redirect(url_for("meetings.manage_conflicts", motion_id=motion.id))
        flash("Select two different amendments", "error")

    conflicts = (
        AmendmentConflict.query.join(
            Amendment, Amendment.id == AmendmentConflict.amendment_a_id
        )
        .filter(Amendment.motion_id == motion.id)
        .all()
    )
    return render_template(
        "meetings/manage_conflicts.html",
        motion=motion,
        amendments=amendments,
        conflicts=conflicts,
        form=form,
    )


@bp.route("/conflicts/<int:conflict_id>/delete", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def delete_conflict(conflict_id: int):
    conflict = db.session.get(AmendmentConflict, conflict_id)
    if conflict is None:
        abort(404)
    motion_id = db.session.get(Amendment, conflict.amendment_a_id).motion_id
    db.session.delete(conflict)
    db.session.commit()
    flash("Conflict removed", "success")
    return redirect(url_for("meetings.manage_conflicts", motion_id=motion_id))


def _amendment_results(meeting: Meeting) -> list[tuple[Amendment, dict]]:
    """Return vote counts for each amendment."""
    amendments = (
        Amendment.query.filter_by(meeting_id=meeting.id).order_by(Amendment.order).all()
    )
    results = []
    for amend in amendments:
        counts = {
            "for": 0,
            "against": 0,
            "abstain": 0,
        }
        rows = (
            db.session.query(Vote.choice, func.count(Vote.id))
            .filter_by(amendment_id=amend.id)
            .filter(Vote.is_test.is_(False))
            .group_by(Vote.choice)
            .all()
        )
        for choice, count in rows:
            counts[choice] = count
        results.append((amend, counts))
    return results


def _motion_results(meeting: Meeting) -> list[tuple[Motion, dict]]:
    """Return vote counts for each motion."""
    motions = (
        Motion.query.filter_by(meeting_id=meeting.id).order_by(Motion.ordering).all()
    )
    results: list[tuple[Motion, dict]] = []
    for motion in motions:
        counts = {
            "for": 0,
            "against": 0,
            "abstain": 0,
        }
        rows = (
            db.session.query(Vote.choice, func.count(Vote.id))
            .filter_by(motion_id=motion.id)
            .filter(Vote.is_test.is_(False))
            .group_by(Vote.choice)
            .all()
        )
        for choice, count in rows:
            counts[choice] = count
        results.append((motion, counts))
    return results


def _merge_form(motions: list[Motion]) -> FlaskForm:
    fields = {}
    for motion in motions:
        fields[f"motion_{motion.id}"] = TextAreaField(
            "Final Motion Text", validators=[DataRequired()]
        )
    fields["submit"] = SubmitField("Save and Send Stage 2 Links")
    return type("MergeForm", (FlaskForm,), fields)()


@bp.route("/<int:meeting_id>/close-stage1", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def close_stage1(meeting_id: int):
    """Close Stage 1 and handle run-offs or open Stage 2."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode == "in-person":
        abort(404)

    meeting.stage1_closed_at = datetime.utcnow()

    # check whether Stage 1 reached quorum
    if meeting.stage1_votes_count() < meeting.quorum:
        meeting.status = "Quorum not met"
        db.session.commit()
        if meeting.opens_at_stage2 and meeting.stage1_closed_at:
            if meeting.opens_at_stage2 - meeting.stage1_closed_at < timedelta(hours=24):
                flash(
                    "Stage 2 opens less than 24 hours after Stage 1 closed",
                    "warning",
                )
        members = Member.query.filter_by(meeting_id=meeting.id).all()
        if auto_send_enabled(meeting, 'quorum_failure'):
            for member in members:
                send_quorum_failure(member, meeting)
        else:
            flash("Automatic emails disabled - use manual send", "warning")
        flash(
            "Stage 1 quorum not met – vote void under Articles 112(d)–(f).",
            "error",
        )
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))

    # if no amendments were proposed, skip Stage 1 entirely
    if Amendment.query.filter_by(meeting_id=meeting.id).count() == 0:
        members = Member.query.filter_by(meeting_id=meeting.id).all()
        if meeting.ballot_mode != "in-person":
            for member in members:
                VoteToken.create(
                    member_id=member.id,
                    stage=2,
                    salt=current_app.config["TOKEN_SALT"],
                )
        meeting.status = "Pending Stage 2"
        db.session.commit()
        if meeting.opens_at_stage2 and meeting.stage1_closed_at:
            if meeting.opens_at_stage2 - meeting.stage1_closed_at < timedelta(hours=24):
                flash(
                    "Stage 2 opens less than 24 hours after Stage 1 closed",
                    "warning",
                )
        flash(
            "No amendments submitted – Stage 1 skipped and Stage 2 tokens generated",
            "success",
        )
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))

    # finalize Stage 1 results and create run-off ballots if required
    runoffs, tokens_to_send = runoff.close_stage1(meeting)

    if runoffs:
        if auto_send_enabled(meeting, 'runoff_invite'):
            for recipient, target, token in tokens_to_send:
                if target:
                    send_proxy_invite(recipient, target, token, meeting)
                else:
                    send_runoff_invite(recipient, token, meeting)
        else:
            flash("Automatic emails disabled - use manual send", "warning")
        flash("Run-off ballot issued; Stage 2 start delayed", "success")
    else:
        meeting.status = "Pending Stage 2"
        db.session.commit()
        if meeting.opens_at_stage2 and meeting.stage1_closed_at:
            if meeting.opens_at_stage2 - meeting.stage1_closed_at < timedelta(hours=24):
                flash(
                    "Stage 2 opens less than 24 hours after Stage 1 closed",
                    "warning",
                )
        flash(
            "Stage 1 closed. Prepare final motion text before opening Stage 2",
            "success",
        )
    return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/results")
@login_required
@permission_required("manage_meetings")
def results_summary(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    stage1_results = _amendment_results(meeting)
    stage2_results: list[tuple[Motion, dict]] = []
    if meeting.status == "Completed":
        stage2_results = _motion_results(meeting)

    stage = 2 if meeting.status in {"Stage 2", "Pending Stage 2"} else 1
    members = (
        Member.query.filter_by(meeting_id=meeting.id, is_test=False)
        .order_by(Member.name)
        .all()
    )
    tokens = (
        VoteToken.query.filter_by(stage=stage, used_at=None, is_test=False)
        .filter(VoteToken.proxy_holder_id.isnot(None))
        .all()
    )
    unused_proxy_tokens = []
    for t in tokens:
        proxy = db.session.get(Member, t.proxy_holder_id)
        principal = db.session.get(Member, t.member_id)
        if (
            proxy is not None
            and principal is not None
            and proxy.meeting_id == meeting.id
            and principal.meeting_id == meeting.id
        ):
            unused_proxy_tokens.append((t, proxy, principal))

    manual_email_mode = AppSetting.get("manual_email_mode") == "1"

    return render_template(
        "meetings/results_summary.html",
        meeting=meeting,
        stage1_results=stage1_results,
        stage2_results=stage2_results,
        members=members,
        unused_proxy_tokens=unused_proxy_tokens,
        manual_email_mode=manual_email_mode,
    )


@bp.route("/<int:meeting_id>/email-settings")
@login_required
@permission_required("manage_meetings")
def email_settings(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    schedule = _email_schedule(meeting)
    settings = {s.email_type: s for s in meeting.email_settings}
    logs = {
        t: EmailLog.query.filter_by(meeting_id=meeting.id, type=t).order_by(EmailLog.sent_at.desc()).first()
        for t in schedule.keys()
    }
    return render_template(
        "meetings/email_settings.html",
        meeting=meeting,
        schedule=schedule,
        settings=settings,
        logs=logs,
    )


@bp.route("/<int:meeting_id>/email-preview/<email_type>")
@login_required
@permission_required("manage_meetings")
def preview_email(meeting_id: int, email_type: str):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    member = (
        Member.query.filter_by(meeting_id=meeting.id).first()
        or Member(name="Example Member", email="test@example.com")
    )
    branding = _branding()
    unsubscribe = "#"
    resubscribe = "#"
    link = url_for("voting.ballot_token", token="preview", _external=True)
    objection_link = url_for(
        "main.public_meeting_detail", meeting_id=meeting.id, _external=True
    )
    if email_type == "stage1_invite":
        notice_html = markdown_to_html(meeting.notice_md or "")
        html = render_template(
            "email/invite.html",
            member=member,
            meeting=meeting,
            link=link,
            notice_html=notice_html,
            objection_link=objection_link,
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    elif email_type == "stage1_reminder":
        template_base = config_or_setting("REMINDER_TEMPLATE", "email/reminder")
        html = render_template(
            f"{template_base}.html",
            member=member,
            meeting=meeting,
            link=link,
            objection_link=objection_link,
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    elif email_type == "stage2_invite":
        summary = carried_amendment_summary(meeting)
        if summary:
            results_link = None
        else:
            if meeting.early_public_results:
                results_link = url_for(
                    "main.public_stage1_results",
                    meeting_id=meeting.id,
                    _external=True,
                )
            else:
                results_link = url_for(
                    "main.public_results", meeting_id=meeting.id, _external=True
                )
        html = render_template(
            "email/stage2_invite.html",
            member=member,
            meeting=meeting,
            link=link,
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            summary=summary,
            results_link=results_link,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    elif email_type == "stage2_reminder":
        template_base = config_or_setting(
            "STAGE2_REMINDER_TEMPLATE", "email/stage2_reminder"
        )
        summary = carried_amendment_summary(meeting)
        if summary:
            results_link = None
        else:
            if meeting.early_public_results:
                results_link = url_for(
                    "main.public_stage1_results",
                    meeting_id=meeting.id,
                    _external=True,
                )
            else:
                results_link = url_for(
                    "main.public_results", meeting_id=meeting.id, _external=True
                )
        html = render_template(
            f"{template_base}.html",
            member=member,
            meeting=meeting,
            link=link,
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            summary=summary,
            results_link=results_link,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    elif email_type == "submission_invite":
        html = render_template(
            "email/submission_invite.html",
            member=member,
            meeting=meeting,
            link=url_for('submissions.submit_motion', token='preview', meeting_id=meeting.id, _external=True),
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    elif email_type == "review_invite":
        html = render_template(
            "email/review_invite.html",
            member=member,
            meeting=meeting,
            review_url=url_for('main.review_motions', token='preview', meeting_id=meeting.id, _external=True),
            link=url_for('submissions.submit_amendment_select', token='preview', meeting_id=meeting.id, _external=True),
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    elif email_type == "amendment_review_invite":
        html = render_template(
            "email/amendment_review_invite.html",
            member=member,
            meeting=meeting,
            review_url=url_for('main.review_motions', token='preview', meeting_id=meeting.id, _external=True),
            unsubscribe_url=unsubscribe,
            resubscribe_url=resubscribe,
            why_text=config_or_setting('EMAIL_WHY_TEXT', 'You are a member of our organisation and have therefore been invited to participate in voting in AGMs/EGMs. If you do not want to participate in the process then please ignore this and subsequent emails'),
            test_mode=True,
            **branding,
        )
    else:
        abort(404)
    return html


@bp.post("/<int:meeting_id>/email-settings/<email_type>")
@login_required
@permission_required("manage_meetings")
def toggle_email_setting(meeting_id: int, email_type: str):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    setting = EmailSetting.query.filter_by(meeting_id=meeting.id, email_type=email_type).first()
    if not setting:
        setting = EmailSetting(meeting_id=meeting.id, email_type=email_type)
        db.session.add(setting)
    setting.auto_send = not setting.auto_send
    db.session.commit()
    schedule = _email_schedule(meeting)
    log = EmailLog.query.filter_by(meeting_id=meeting.id, type=email_type).order_by(EmailLog.sent_at.desc()).first()
    return render_template(
        "meetings/_email_row.html",
        meeting=meeting,
        email_type=email_type,
        schedule_time=schedule.get(email_type),
        setting=setting,
        log=log,
    )


@bp.route("/<int:meeting_id>/send-emails", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def manual_send_emails(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode == "in-person":
        abort(404)
    form = ManualEmailForm()
    members = (
        Member.query.filter_by(meeting_id=meeting.id, is_test=False)
        .order_by(Member.name)
        .all()
    )
    form.member_ids.choices = [(m.id, m.name) for m in members]
    if form.validate_on_submit():
        recipients = []
        if form.test_mode.data:
            current_user_member = Member.query.filter_by(
                meeting_id=meeting.id, email=current_user.email
            ).first()
            if not current_user_member:
                current_user_member = Member(
                    meeting_id=meeting.id,
                    name=current_user.email,
                    email=current_user.email,
                    is_test=True,
                )
                db.session.add(current_user_member)
                db.session.commit()
            recipients = [current_user_member]
        elif form.send_to_all.data:
            recipients = members
        else:
            recipients = [m for m in members if m.id in form.member_ids.data]

        for member in recipients:
            stage = 1
            if form.email_type.data == "stage2_invite":
                stage = 2
            token_obj, plain = VoteToken.create(
                member_id=member.id,
                stage=stage,
                salt=current_app.config["TOKEN_SALT"],
            )
            token_obj.is_test = form.test_mode.data
            db.session.commit()

            if form.email_type.data == "stage1_invite":
                send_vote_invite(member, plain, meeting, test_mode=form.test_mode.data)
            elif form.email_type.data == "stage1_reminder":
                send_stage1_reminder(
                    member, plain, meeting, test_mode=form.test_mode.data
                )
            elif form.email_type.data == "runoff_invite":
                send_runoff_invite(
                    member, plain, meeting, test_mode=form.test_mode.data
                )
            elif form.email_type.data == "stage2_invite":
                send_stage2_invite(
                    member, plain, meeting, test_mode=form.test_mode.data
                )
            elif form.email_type.data == "submission_invite":
                send_submission_invite(member, meeting, test_mode=form.test_mode.data)
            elif form.email_type.data == "review_invite":
                send_review_invite(member, meeting, test_mode=form.test_mode.data)
            elif form.email_type.data == "amendment_review_invite":
                send_amendment_review_invite(member, meeting, test_mode=form.test_mode.data)

        flash("Emails sent", "success")
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))

    return render_template("meetings/manual_email.html", meeting=meeting, form=form)


@bp.route("/<int:meeting_id>/extend/<int:stage>", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def extend_stage(meeting_id: int, stage: int):
    """Allow admin to extend Stage 1 or 2 voting window."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or stage not in (1, 2):
        abort(404)
    form = ExtendStageForm()
    if request.method == "GET":
        if stage == 1:
            form.opens_at.data = meeting.opens_at_stage1
            form.closes_at.data = meeting.closes_at_stage1
        else:
            form.opens_at.data = meeting.opens_at_stage2
            form.closes_at.data = meeting.closes_at_stage2
    if form.validate_on_submit():
        if stage == 1:
            meeting.opens_at_stage1 = form.opens_at.data
            meeting.closes_at_stage1 = form.closes_at.data
        else:
            meeting.opens_at_stage2 = form.opens_at.data
            meeting.closes_at_stage2 = form.closes_at.data
        meeting.extension_reason = form.reason.data
        db.session.commit()
        flash("Stage dates updated", "success")
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))
    return render_template(
        "meetings/extend_stage.html",
        form=form,
        meeting=meeting,
        stage=stage,
    )


@bp.route("/<int:meeting_id>/stage1-tally", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def stage1_tally(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode != "in-person":
        abort(404)
    form = Stage1TallyForm()
    if request.method == "GET":
        form.votes_cast.data = meeting.stage1_manual_votes
    if form.validate_on_submit():
        meeting.stage1_manual_votes = form.votes_cast.data or 0
        meeting.status = "Pending Stage 2"
        db.session.commit()
        flash("Stage 1 tally saved", "success")
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))
    return render_template(
        "meetings/stage1_tally_form.html",
        form=form,
        meeting=meeting,
    )


@bp.route("/<int:meeting_id>/stage2-tally", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def stage2_tally(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode != "in-person":
        abort(404)
    form = Stage2TallyForm()
    if request.method == "GET":
        form.for_votes.data = meeting.stage2_manual_for
        form.against_votes.data = meeting.stage2_manual_against
        form.abstain_votes.data = meeting.stage2_manual_abstain
    if form.validate_on_submit():
        meeting.stage2_manual_for = form.for_votes.data or 0
        meeting.stage2_manual_against = form.against_votes.data or 0
        meeting.stage2_manual_abstain = form.abstain_votes.data or 0
        meeting.status = "Completed"
        db.session.commit()
        flash("Stage 2 tally saved", "success")
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))
    return render_template(
        "meetings/stage2_tally_form.html",
        form=form,
        meeting=meeting,
    )


@bp.route("/<int:meeting_id>/preview/comments/motion/<int:motion_id>", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def preview_motion_comments(meeting_id: int, motion_id: int):
    if request.method == "POST":
        return comments.add_motion_comment("preview", motion_id)
    return comments.motion_comments("preview", motion_id)


@bp.route("/<int:meeting_id>/preview/comments/amendment/<int:amendment_id>", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def preview_amendment_comments(meeting_id: int, amendment_id: int):
    if request.method == "POST":
        return comments.add_amendment_comment("preview", amendment_id)
    return comments.amendment_comments("preview", amendment_id)


@bp.route("/<int:meeting_id>/preview/<int:stage>", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def preview_voting(meeting_id: int, stage: int):
    """Display a preview of the voting screens without recording votes."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)

    if meeting.ballot_mode == "combined":
        motions = (
            Motion.query.filter_by(meeting_id=meeting.id)
            .order_by(Motion.ordering)
            .all()
        )
        amendments = (
            Amendment.query.filter(Amendment.motion_id.in_([m.id for m in motions]))
            .order_by(Amendment.order)
            .all()
        )
        form = _combined_form(motions, amendments)
        if form.validate_on_submit():
            flash("Preview submission complete – votes were not saved", "info")
            return render_template(
                "voting/confirmation.html",
                preview=True,
                stage=stage,
                final_message_default=current_app.config.get("FINAL_STAGE_MESSAGE"),
            )
        motions_grouped = []
        for motion in motions:
            ams = [a for a in amendments if a.motion_id == motion.id]
            motions_grouped.append((motion, ams))
        motion_counts = {
            m.id: Comment.query.filter_by(motion_id=m.id, hidden=False).count()
            for m in motions
        }
        amend_counts = {
            a.id: Comment.query.filter_by(amendment_id=a.id, hidden=False).count()
            for a in amendments
        }
        return render_template(
            "voting/combined_ballot.html",
            form=form,
            motions=motions_grouped,
            meeting=meeting,
            proxy_for=None,
            token="preview",
            preview=True,
            motion_counts=motion_counts,
            amend_counts=amend_counts,
        )

    if stage == 1:
        motions = (
            Motion.query.filter_by(meeting_id=meeting.id)
            .order_by(Motion.ordering)
            .all()
        )
        amendments = (
            Amendment.query.filter(Amendment.motion_id.in_([m.id for m in motions]))
            .order_by(Amendment.order)
            .all()
        )
        form = _amendment_form(amendments)
        if form.validate_on_submit():
            flash("Preview submission complete – votes were not saved", "info")
            return render_template(
                "voting/confirmation.html",
                preview=True,
                stage=stage,
                final_message_default=current_app.config.get("FINAL_STAGE_MESSAGE"),
            )
        motions_grouped = []
        for motion in motions:
            ams = [a for a in amendments if a.motion_id == motion.id]
            motions_grouped.append((motion, ams))
        motion_counts = {
            m.id: Comment.query.filter_by(motion_id=m.id, hidden=False).count()
            for m in motions
        }
        amend_counts = {
            a.id: Comment.query.filter_by(amendment_id=a.id, hidden=False).count()
            for a in amendments
        }
        return render_template(
            "voting/stage1_ballot.html",
            form=form,
            motions=motions_grouped,
            meeting=meeting,
            proxy_for=None,
            token="preview",
            preview=True,
            motion_counts=motion_counts,
            amend_counts=amend_counts,
        )

    motions = (
        Motion.query.filter_by(meeting_id=meeting.id).order_by(Motion.ordering).all()
    )
    form = _motion_form(motions)
    if form.validate_on_submit():
        flash("Preview submission complete – votes were not saved", "info")
        return render_template(
            "voting/confirmation.html",
            preview=True,
            stage=stage,
            final_message_default=current_app.config.get("FINAL_STAGE_MESSAGE"),
        )
    compiled = [(m, m.final_text_md or compile_motion_text(m)) for m in motions]
    motion_counts = {
        m.id: Comment.query.filter_by(motion_id=m.id, hidden=False).count()
        for m in motions
    }
    return render_template(
        "voting/stage2_ballot.html",
        form=form,
        motions=compiled,
        meeting=meeting,
        proxy_for=None,
        token="preview",
        preview=True,
        motion_counts=motion_counts,
    )


@bp.route("/<int:meeting_id>/results.docx")
@login_required
@permission_required("manage_meetings")
def results_docx(meeting_id: int):
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    results = _amendment_results(meeting)
    include_logo = request.args.get("logo") == "1"
    doc = _styled_doc(f"{meeting.title} - Stage 1 Results", include_logo)

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Amendment"
    hdr[1].text = "For"
    hdr[2].text = "Against"
    hdr[3].text = "Abstain"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for idx, (amend, counts) in enumerate(results, start=1):
        row = table.add_row().cells
        row[0].text = amend.text_md or ""
        row[1].text = str(counts["for"])
        row[2].text = str(counts["against"])
        row[3].text = str(counts["abstain"])
        if idx % 2 == 0:
            for c in row:
                _shade_cell(c, "F7F7F9")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="results.docx",
    )
    resp.headers["Content-Disposition"] = 'attachment; filename="results.docx"'
    return resp


@bp.route("/<int:meeting_id>/prepare-stage2", methods=["GET", "POST"])
@login_required
@permission_required("manage_meetings")
def prepare_stage2(meeting_id: int):
    """Allow a human to merge amendments into final motion text."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode == "in-person":
        abort(404)
    if meeting.opens_at_stage2 and meeting.stage1_closed_at:
        if meeting.opens_at_stage2 - meeting.stage1_closed_at < timedelta(hours=24):
            flash(
                "Stage 2 opens less than 24 hours after Stage 1 closed",
                "warning",
            )
    motions = (
        Motion.query.filter_by(meeting_id=meeting.id).order_by(Motion.ordering).all()
    )
    form = _merge_form(motions)
    if form.validate_on_submit():
        for motion in motions:
            data = form[f"motion_{motion.id}"].data
            motion.final_text_md = data
        members = Member.query.filter_by(meeting_id=meeting.id).all()
        tokens_to_send: list[tuple[Member, str]] = []
        proxy_tokens: list[tuple[Member, Member, str]] = []
        for member in members:
            token_obj, plain = VoteToken.create(
                member_id=member.id,
                stage=2,
                salt=current_app.config["TOKEN_SALT"],
            )
            tokens_to_send.append((member, plain))
        for proxy in members:
            if proxy.proxy_for:
                try:
                    target = db.session.get(Member, int(proxy.proxy_for))
                except (ValueError, TypeError):
                    target = None
                if target:
                    token_obj, plain = VoteToken.create(
                        member_id=target.id,
                        stage=2,
                        salt=current_app.config["TOKEN_SALT"],
                        proxy_holder_id=proxy.id,
                    )
                    proxy_tokens.append((proxy, target, plain))
        meeting.status = "Stage 2"
        db.session.commit()
        if auto_send_enabled(meeting, 'stage2_invite'):
            for m, t in tokens_to_send:
                send_stage2_invite(m, t, meeting)
            for p, target, tok in proxy_tokens:
                send_proxy_invite(p, target, tok, meeting)
        else:
            flash("Automatic emails disabled - use manual send", "warning")
        flash("Stage 2 voting links sent", "success")
        return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))

    for motion in motions:
        field = form[f"motion_{motion.id}"]
        field.data = motion.final_text_md or compile_motion_text(motion)

    return render_template(
        "meetings/prepare_stage2.html",
        meeting=meeting,
        motions=motions,
        form=form,
        compile_motion_text=compile_motion_text,
    )


@bp.route("/<int:meeting_id>/results-stage2.docx")
def results_stage2_docx(meeting_id: int):
    """Download DOCX summarising carried amendments and final motion results."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    if not meeting.public_results or not meeting.results_doc_published:
        abort(404)
    amend_results = _amendment_results(meeting)
    motion_results = _motion_results(meeting)

    include_logo = request.args.get("logo") == "1"
    doc = _styled_doc(f"{meeting.title} - Final Results", include_logo)
    para = doc.add_paragraph(
        "This document is a draft summary. The organisation reserves the right to issue a final official version."
    )
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if meeting.results_doc_intro_md:
        doc.add_paragraph(meeting.results_doc_intro_md)

    doc.add_heading("Carried Amendments", level=2)
    carried = [a for a, c in amend_results if c.get("for", 0) > c.get("against", 0)]
    table_ca = doc.add_table(rows=1, cols=1)
    if carried:
        for idx, amend in enumerate(carried, start=1):
            row = table_ca.add_row().cells
            row[0].text = amend.text_md or ""
            if idx % 2 == 0:
                _shade_cell(row[0], "F7F7F9")
    else:
        table_ca.rows[0].cells[0].text = "No amendments carried."

    doc.add_heading("Motion Outcomes", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Motion"
    hdr[1].text = "For"
    hdr[2].text = "Against"
    hdr[3].text = "Abstain"
    hdr[4].text = "Outcome"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for idx, (motion, counts) in enumerate(motion_results, start=1):
        row = table.add_row().cells
        row[0].text = motion.title or "Motion"
        row[1].text = str(counts["for"])
        row[2].text = str(counts["against"])
        row[3].text = str(counts["abstain"])
        row[4].text = (motion.status or "?").capitalize()
        if idx % 2 == 0:
            for c in row:
                _shade_cell(c, "F7F7F9")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="final_results.docx",
    )
    resp.headers["Content-Disposition"] = 'attachment; filename="final_results.docx"'
    return resp


@bp.route("/<int:meeting_id>/stage1.ics")
@login_required
@permission_required("manage_meetings")
def stage1_ics(meeting_id: int):
    """Download Stage 1 calendar file."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    try:
        ics = generate_stage_ics(meeting, 1)
    except ValueError:
        flash("Stage 1 timestamps not set", "error")
        return redirect(request.referrer or url_for("meetings.list_meetings"))
    return send_file(
        io.BytesIO(ics),
        mimetype="text/calendar",
        as_attachment=True,
        download_name="stage1.ics",
    )


@bp.route("/<int:meeting_id>/stage2.ics")
@login_required
@permission_required("manage_meetings")
def stage2_ics(meeting_id: int):
    """Download Stage 2 calendar file."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    try:
        ics = generate_stage_ics(meeting, 2)
    except ValueError:
        flash("Stage 2 timestamps not set", "error")
        return redirect(request.referrer or url_for("meetings.list_meetings"))
    return send_file(
        io.BytesIO(ics),
        mimetype="text/calendar",
        as_attachment=True,
        download_name="stage2.ics",
    )


@bp.route("/<int:meeting_id>/close-runoff", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def close_runoff(meeting_id: int):
    """Finalize run-off results and move meeting to Pending Stage 2."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    runoff.close_runoff_stage(meeting)
    meeting.status = "Pending Stage 2"
    db.session.commit()
    flash("Run-off votes tallied", "success")
    return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/close-stage2", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def close_stage2(meeting_id: int):
    """Finalize Stage 2 results and record motion outcomes."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode == "in-person":
        abort(404)
    motion_results = _motion_results(meeting)

    for motion, counts in motion_results:
        total = counts["for"] + counts["against"] + counts["abstain"]
        if total == 0:
            motion.status = "failed"
            continue
        ratio = counts["for"] / total
        if motion.threshold == "special":
            carried = ratio >= 0.75
        else:
            carried = counts["for"] > counts["against"]
        motion.status = "carried" if carried else "failed"

    meeting.status = "Completed"
    db.session.commit()

    members = Member.query.filter_by(meeting_id=meeting.id).all()
    if auto_send_enabled(meeting, 'final_results'):
        for member in members:
            send_final_results(member, meeting)
    else:
        flash("Automatic emails disabled - use manual send", "warning")

    flash("Stage 2 closed and motions tallied", "success")
    return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/members/<int:member_id>/resend", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def resend_member_link(meeting_id: int, member_id: int):
    """Generate a new voting token for the current stage and email it."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode == "in-person":
        abort(404)
    member = Member.query.filter_by(id=member_id, meeting_id=meeting.id).first_or_404()

    stage = 2 if meeting.status in {"Stage 2", "Pending Stage 2"} else 1

    token_obj, plain = VoteToken.create(
        member_id=member.id, stage=stage, salt=current_app.config["TOKEN_SALT"]
    )
    db.session.commit()

    if stage == 2:
        send_stage2_invite(member, plain, meeting)
    else:
        if Runoff.query.filter_by(meeting_id=meeting.id).count() > 0:
            send_runoff_invite(member, plain, meeting)
        else:
            send_vote_invite(member, plain, meeting)

    flash("Voting link resent", "success")
    return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/members/<int:member_id>/email/<kind>", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def send_member_email(meeting_id: int, member_id: int, kind: str):
    """Send a specific meeting email to a member if allowed."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.ballot_mode == "in-person":
        abort(404)
    member = Member.query.filter_by(id=member_id, meeting_id=meeting.id).first_or_404()

    now = datetime.utcnow()

    def within(start: datetime | None, end: datetime | None) -> bool:
        return start is not None and start <= now and (end is None or now <= end)

    if kind == "stage1_invite":
        if not within(meeting.opens_at_stage1, meeting.opens_at_stage2):
            abort(400)
        token, plain = VoteToken.create(
            member_id=member.id, stage=1, salt=current_app.config["TOKEN_SALT"]
        )
        db.session.commit()
        send_vote_invite(member, plain, meeting)
        flash("Stage 1 invite sent", "success")
    elif kind == "stage1_reminder":
        if not within(meeting.opens_at_stage1, meeting.closes_at_stage1):
            abort(400)
        token, plain = VoteToken.create(
            member_id=member.id, stage=1, salt=current_app.config["TOKEN_SALT"]
        )
        db.session.commit()
        send_stage1_reminder(member, plain, meeting)
        flash("Stage 1 reminder sent", "success")
    elif kind == "runoff_invite":
        if (
            Runoff.query.filter_by(meeting_id=meeting.id).count() == 0
            or meeting.opens_at_stage2 is None
            or not meeting.stage1_closed_at
            or not (meeting.stage1_closed_at <= now < meeting.opens_at_stage2)
        ):
            abort(400)
        token, plain = VoteToken.create(
            member_id=member.id, stage=1, salt=current_app.config["TOKEN_SALT"]
        )
        db.session.commit()
        send_runoff_invite(member, plain, meeting)
        flash("Run-off invite sent", "success")
    elif kind == "stage2_invite":
        if meeting.ballot_mode != "two-stage" or not within(
            meeting.opens_at_stage2, meeting.closes_at_stage2
        ):
            abort(400)
        token, plain = VoteToken.create(
            member_id=member.id, stage=2, salt=current_app.config["TOKEN_SALT"]
        )
        db.session.commit()
        send_stage2_invite(member, plain, meeting)
        flash("Stage 2 invite sent", "success")
    elif kind == "stage2_reminder":
        if meeting.ballot_mode != "two-stage" or not within(
            meeting.opens_at_stage2, meeting.closes_at_stage2
        ):
            abort(400)
        token, plain = VoteToken.create(
            member_id=member.id, stage=2, salt=current_app.config["TOKEN_SALT"]
        )
        db.session.commit()
        send_stage2_reminder(member, plain, meeting)
        flash("Stage 2 reminder sent", "success")
    elif kind == "submission_invite":
        if not within(meeting.motions_opens_at, meeting.motions_closes_at):
            abort(400)
        send_submission_invite(member, meeting)
        flash("Submission invite sent", "success")
    elif kind == "review_invite":
        if not within(meeting.amendments_opens_at, meeting.amendments_closes_at):
            abort(400)
        send_review_invite(member, meeting)
        flash("Review invite sent", "success")
    elif kind == "amendment_review_invite":
        if not within(meeting.amendments_closes_at, meeting.opens_at_stage1):
            abort(400)
        send_amendment_review_invite(member, meeting)
        flash("Amendment review invite sent", "success")
    elif kind == "final_results":
        if meeting.status != "Completed":
            abort(400)
        send_final_results(member, meeting)
        flash("Final results sent", "success")
    else:
        abort(404)

    return redirect(url_for("meetings.list_members", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/send-final-results", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def send_final_results_all(meeting_id: int):
    """Email certified results to all meeting members."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None or meeting.status != "Completed":
        abort(404)

    members = Member.query.filter_by(meeting_id=meeting.id).all()
    for member in members:
        send_final_results(member, meeting)

    flash("Final results emailed", "success")
    return redirect(url_for("meetings.meeting_overview", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/proxy-tokens/<token>/resend", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def resend_proxy_token(meeting_id: int, token: str):
    """Generate a new proxy token and email it."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    vt = db.session.get(VoteToken, token)
    if vt is None or vt.proxy_holder_id is None:
        abort(404)
    proxy = Member.query.filter_by(
        id=vt.proxy_holder_id, meeting_id=meeting.id
    ).first_or_404()
    principal = Member.query.filter_by(
        id=vt.member_id, meeting_id=meeting.id
    ).first_or_404()

    new_token, plain = VoteToken.create(
        member_id=principal.id,
        stage=vt.stage,
        salt=current_app.config["TOKEN_SALT"],
        proxy_holder_id=proxy.id,
    )
    db.session.commit()

    send_proxy_invite(proxy, principal, plain, meeting)
    flash("Proxy voting link resent", "success")
    return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))


@bp.route("/<int:meeting_id>/proxy-tokens/<token>/invalidate", methods=["POST"])
@login_required
@permission_required("manage_meetings")
def invalidate_proxy_token(meeting_id: int, token: str):
    """Mark a proxy token as used to invalidate it."""
    meeting = db.session.get(Meeting, meeting_id)
    if meeting is None:
        abort(404)
    vt = db.session.get(VoteToken, token)
    if vt is None or vt.proxy_holder_id is None:
        abort(404)
    proxy = db.session.get(Member, vt.proxy_holder_id)
    principal = db.session.get(Member, vt.member_id)
    if (
        proxy is None
        or principal is None
        or proxy.meeting_id != meeting.id
        or principal.meeting_id != meeting.id
    ):
        abort(404)

    vt.used_at = datetime.utcnow()
    db.session.commit()
    flash("Proxy token invalidated", "success")
    return redirect(url_for("meetings.results_summary", meeting_id=meeting.id))
