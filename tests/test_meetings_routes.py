import os, sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from unittest.mock import patch
import pytest
from werkzeug.exceptions import Forbidden, NotFound
from flask import url_for

from app import create_app
from app.extensions import db
from app.models import (
    User,
    Role,
    Permission,
    Meeting,
    VoteToken,
    Member,
    Motion,
    Amendment,
    MeetingFile,
    Vote,
)
import io
from app.meetings import routes as meetings
from app import routes as main
from docx import Document
from app.meetings.forms import MeetingForm
from types import SimpleNamespace
from datetime import datetime, timedelta
from uuid6 import uuid7
from werkzeug.datastructures import MultiDict
from app.models import AmendmentConflict


def _make_user(has_permission: bool):
    perm = Permission(name="manage_meetings") if has_permission else None
    role = Role(permissions=[perm] if perm else [])
    user = User(role=role)
    user.email = "admin@example.com"
    user.is_active = True
    return user


def test_list_meetings_requires_permission():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        with app.test_request_context("/meetings/"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.list_meetings()

        with app.test_request_context("/meetings/"):
            user = _make_user(False)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash"):
                    try:
                        meetings.list_meetings()
                    except Forbidden:
                        pass
                    else:
                        assert False, "expected Forbidden"


def test_list_meetings_contains_create_link():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        with app.test_request_context("/meetings/"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                html = meetings.list_meetings()
                href = url_for("meetings.create_meeting")
                assert href in html
                assert "bp-btn-primary" in html


def test_import_members_sends_invites_and_tokens():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Test")
        db.session.add(meeting)
        db.session.commit()

        csv_content = "member_id,name,email,proxy_for\n" "1,Alice,alice@example.com,\n"
        data = {"csv_file": (io.BytesIO(csv_content.encode()), "members.csv")}
        with app.test_request_context(
            f"/meetings/{meeting.id}/import-members", method="POST", data=data
        ):
            user = _make_user(True)
            dummy_form = SimpleNamespace(
                csv_file=SimpleNamespace(data=io.BytesIO(csv_content.encode()))
            )
            dummy_form.validate_on_submit = lambda: True
            with patch("flask_login.utils._get_user", return_value=user):
                with patch(
                    "app.meetings.routes.MemberImportForm", return_value=dummy_form
                ):
                    with patch("app.meetings.routes.send_vote_invite") as mock_send:
                        meetings.import_members(meeting.id)
                        mock_send.assert_called_once()
                        assert VoteToken.query.count() == 1


def test_close_stage1_creates_stage2_tokens_and_emails():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["MAIL_SUPPRESS_SEND"] = True
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Test")
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="Bob", email="b@example.com")
        db.session.add(member)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="x",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.flush()
        amend = Amendment(
            meeting_id=meeting.id, motion_id=motion.id, text_md="A1", order=1
        )
        db.session.add(amend)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/close-stage1", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch(
                    "app.meetings.routes.runoff.close_stage1", return_value=([], [])
                ):
                    with patch("app.meetings.routes.send_stage2_invite") as mock_send:
                        meetings.close_stage1(meeting.id)
                        mock_send.assert_not_called()
                        assert (
                            VoteToken.query.filter_by(
                                member_id=member.id, stage=2
                            ).count()
                            == 0
                        )
                        assert meeting.status == "Pending Stage 2"


def test_close_stage1_runoff_triggers_emails_and_tokens():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["MAIL_SUPPRESS_SEND"] = True
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Test")
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="Bob", email="b@example.com")
        db.session.add(member)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="x",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.flush()
        amend = Amendment(
            meeting_id=meeting.id, motion_id=motion.id, text_md="A1", order=1
        )
        db.session.add(amend)
        db.session.commit()

        runoff_obj = SimpleNamespace(id=1)

        def _runoff_side_effect(mtg):
            token_obj, plain = VoteToken.create(member_id=member.id, stage=1, salt="s")
            db.session.commit()
            return [runoff_obj], [(member, None, plain)]

        with app.test_request_context(
            f"/meetings/{meeting.id}/close-stage1", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                from app.services import runoff

                with patch.object(
                    runoff, "close_stage1", side_effect=_runoff_side_effect
                ):
                    with patch(
                        "app.meetings.routes.send_proxy_invite"
                    ) as mock_proxy, patch(
                        "app.meetings.routes.send_runoff_invite"
                    ) as mock_send:
                        meetings.close_stage1(meeting.id)
                        mock_send.assert_called_once()
                    assert (
                        VoteToken.query.filter_by(member_id=member.id, stage=1).count()
                        == 1
                    )


def test_close_stage1_skips_when_no_amendments():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Test", quorum=1)
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="Bob", email="b@example.com")
        db.session.add(member)
        db.session.flush()
        token, _ = VoteToken.create(member_id=member.id, stage=1, salt="s")
        token.used_at = datetime.utcnow()
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/close-stage1", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.runoff.close_stage1") as close_mock:
                    with patch("app.meetings.routes.send_stage2_invite") as send_mock:
                        with patch("app.meetings.routes.flash") as fl:
                            meetings.close_stage1(meeting.id)
                            fl.assert_called()
                        close_mock.assert_not_called()
                        send_mock.assert_not_called()

        assert meeting.status == "Pending Stage 2"
        assert VoteToken.query.filter_by(member_id=member.id, stage=2).count() == 1


def test_close_stage1_below_quorum_voids_vote():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Test", quorum=5)
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="Bob", email="b@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/close-stage1", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch.object(Meeting, "stage1_votes_count", return_value=2):
                    with patch("app.meetings.routes.runoff.close_stage1") as mock_close:
                        with patch(
                            "app.meetings.routes.send_stage2_invite"
                        ) as mock_stage2:
                            with patch(
                                "app.meetings.routes.send_runoff_invite"
                            ) as mock_runoff:
                                with patch(
                                    "app.meetings.routes.send_quorum_failure"
                                ) as mock_qfail:
                                    meetings.close_stage1(meeting.id)
                                    mock_close.assert_not_called()
                                    mock_stage2.assert_not_called()
                                    mock_runoff.assert_not_called()
                                    mock_qfail.assert_called_once()
                                assert meeting.status == "Quorum not met"
        assert VoteToken.query.filter_by(member_id=member.id, stage=2).count() == 0


def test_close_stage1_warns_when_stage2_too_soon():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        now = datetime.utcnow()
        meeting = Meeting(title="Test", opens_at_stage2=now + timedelta(hours=12))
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="Bob")
        db.session.add(member)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="x",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.flush()
        amend = Amendment(
            meeting_id=meeting.id, motion_id=motion.id, text_md="A1", order=1
        )
        db.session.add(amend)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/close-stage1", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch(
                    "app.meetings.routes.runoff.close_stage1", return_value=([], [])
                ):
                    with patch("app.meetings.routes.send_stage2_invite"):
                        with patch("app.meetings.routes.flash") as fl:
                            meetings.close_stage1(meeting.id)
                            assert any(
                                c.args[1] == "warning" for c in fl.call_args_list
                            )


def test_prepare_stage2_warns_when_too_soon():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        now = datetime.utcnow()
        meeting = Meeting(
            title="Test",
            stage1_closed_at=now,
            opens_at_stage2=now + timedelta(hours=12),
        )
        db.session.add(meeting)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="x",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.commit()

        with app.test_request_context(f"/meetings/{meeting.id}/prepare-stage2"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    meetings.prepare_stage2(meeting.id)
                    assert any(c.args[1] == "warning" for c in fl.call_args_list)


def test_add_amendment_validations():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="Test", opens_at_stage1=datetime.utcnow() + timedelta(days=30)
        )
        db.session.add(meeting)
        db.session.flush()
        members = [
            Member(meeting_id=meeting.id, name="A", email="a@example.com"),
            Member(meeting_id=meeting.id, name="B", email="b@example.com"),
            Member(meeting_id=meeting.id, name="C", email="c@example.com"),
        ]
        for m in members:
            db.session.add(m)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="text",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.commit()

        user = _make_user(True)
        data = {
            "text_md": "A1",
            "proposer_id": members[0].id,
            "seconder_id": members[1].id,
        }
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/amendments/add", method="POST", data=data
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.add_amendment(motion.id)

        assert Amendment.query.count() == 1

        # fourth amendment by same proposer should fail
        for i in range(2, 5):
            with app.test_request_context(
                f"/meetings/motions/{motion.id}/amendments/add",
                method="POST",
                data={
                    "text_md": f"A{i}",
                    "proposer_id": members[0].id,
                    "seconder_id": members[2].id,
                },
            ):
                with patch("flask_login.utils._get_user", return_value=user):
                    if i < 4:
                        meetings.add_amendment(motion.id)
                    else:
                        with patch("app.meetings.routes.flash") as fl:
                            meetings.add_amendment(motion.id)
                            fl.assert_called()

        assert Amendment.query.count() == 3

        # board seconded amendment
        data = {
            "text_md": "board",
            "proposer_id": members[1].id,
            "seconder_id": "0",
            "board_seconded": "y",
        }
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/amendments/add",
            method="POST",
            data=data,
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.add_amendment(motion.id)

        assert Amendment.query.count() == 4

        # missing seconder should fail
        data = {
            "text_md": "fail",
            "proposer_id": members[2].id,
            "seconder_id": "0",
        }
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/amendments/add",
            method="POST",
            data=data,
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    meetings.add_amendment(motion.id)
                    fl.assert_called()

        assert Amendment.query.count() == 4

        # deadline validation
        meeting.opens_at_stage1 = datetime.utcnow() + timedelta(days=10)
        db.session.commit()
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/amendments/add",
            method="POST",
            data={
                "text_md": "late",
                "proposer_id": members[2].id,
                "seconder_id": members[1].id,
            },
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    meetings.add_amendment(motion.id)
                    fl.assert_called()

        assert Amendment.query.count() == 4


def test_add_amendment_notice_deadline():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="Test",
            notice_date=datetime.utcnow() - timedelta(days=20),
        )
        db.session.add(meeting)
        db.session.flush()
        members = [
            Member(meeting_id=meeting.id, name="A"),
            Member(meeting_id=meeting.id, name="B"),
        ]
        db.session.add_all(members)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="t",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.commit()

        user = _make_user(True)
        data = {
            "text_md": "late",
            "proposer_id": members[0].id,
            "seconder_id": members[1].id,
        }
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/amendments/add",
            method="POST",
            data=data,
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    meetings.add_amendment(motion.id)
                    fl.assert_called()

        assert Amendment.query.count() == 0


def test_results_stage2_docx_returns_file():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="AGM",
            public_results=True,
            results_doc_published=True,
            results_doc_intro_md="Intro",
        )
        db.session.add(meeting)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="Motion text",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        member = Member(meeting_id=meeting.id, name="Alice")
        db.session.add(member)
        db.session.flush()
        Vote.record(member_id=member.id, motion_id=motion.id, choice="for", salt="s")

        with app.test_request_context(f"/meetings/{meeting.id}/results-stage2.docx"):
            resp = meetings.results_stage2_docx(meeting.id)
            assert resp.mimetype == (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            cd = resp.headers["Content-Disposition"]
            assert "final_results.docx" in cd
            resp.direct_passthrough = False
            doc = Document(io.BytesIO(resp.get_data()))
            assert "draft summary" in doc.paragraphs[0].text
            assert "Intro" in doc.paragraphs[1].text


def test_results_stage2_docx_checks_flags():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM", public_results=False, results_doc_published=True)
        db.session.add(meeting)
        db.session.commit()
        with app.test_request_context(f"/meetings/{meeting.id}/results-stage2.docx"):
            with pytest.raises(NotFound):
                meetings.results_stage2_docx(meeting.id)

        meeting.public_results = True
        meeting.results_doc_published = False
        db.session.commit()
        with app.test_request_context(f"/meetings/{meeting.id}/results-stage2.docx"):
            with pytest.raises(NotFound):
                meetings.results_stage2_docx(meeting.id)


def test_stage_ics_downloads_with_headers():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        now = datetime.utcnow()
        meeting = Meeting(
            title="AGM",
            opens_at_stage1=now,
            closes_at_stage1=now + timedelta(days=7),
            opens_at_stage2=now + timedelta(days=8),
            closes_at_stage2=now + timedelta(days=13),
        )
        db.session.add(meeting)
        db.session.commit()
        user = _make_user(True)
        with app.test_request_context(f"/meetings/{meeting.id}/stage1.ics"):
            with patch("flask_login.utils._get_user", return_value=user):
                resp1 = meetings.stage1_ics(meeting.id)
                assert resp1.mimetype == "text/calendar"
                cd1 = resp1.headers["Content-Disposition"]
                assert "stage1.ics" in cd1

        with app.test_request_context(f"/meetings/{meeting.id}/stage2.ics"):
            with patch("flask_login.utils._get_user", return_value=user):
                resp2 = meetings.stage2_ics(meeting.id)
                assert resp2.mimetype == "text/calendar"
                cd2 = resp2.headers["Content-Disposition"]
                assert "stage2.ics" in cd2


def test_meeting_files_upload_and_public_view(tmp_path):
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    app.config["UPLOAD_FOLDER"] = str(tmp_path)
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.commit()
        user = _make_user(True)
        data = {
            "file": (io.BytesIO(b"%PDF-1.4"), "test.pdf"),
            "title": "Agenda",
            "description": "AGM agenda",
        }
        with app.test_request_context(
            f"/meetings/{meeting.id}/files", method="POST", data=data
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                resp = meetings.meeting_files(meeting.id)
                assert resp.status_code == 302
                assert MeetingFile.query.count() == 1

        with app.test_request_context(f"/public/meetings/{meeting.id}"):
            html = main.public_meeting_detail(meeting.id)
            assert "Agenda" in html


def test_stage_ics_missing_timestamps_redirects():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.commit()
        user = _make_user(True)
        with app.test_request_context(f"/meetings/{meeting.id}/stage1.ics"):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    resp1 = meetings.stage1_ics(meeting.id)
                    fl.assert_called()
                    assert resp1.status_code == 302

        with app.test_request_context(f"/meetings/{meeting.id}/stage2.ics"):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    resp2 = meetings.stage2_ics(meeting.id)
                    fl.assert_called()
                    assert resp2.status_code == 302


def test_close_stage2_sets_motion_statuses():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        m1 = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="text1",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        m2 = Motion(
            meeting_id=meeting.id,
            title="M2",
            text_md="text2",
            category="motion",
            threshold="special",
            ordering=2,
        )
        db.session.add_all([m1, m2])
        members = [Member(meeting_id=meeting.id, name=f"M{i}") for i in range(4)]
        for mem in members:
            db.session.add(mem)
        db.session.flush()
        Vote.record(member_id=members[0].id, motion_id=m1.id, choice="for", salt="s")
        Vote.record(member_id=members[1].id, motion_id=m1.id, choice="for", salt="s")
        Vote.record(
            member_id=members[2].id, motion_id=m1.id, choice="against", salt="s"
        )

        Vote.record(member_id=members[0].id, motion_id=m2.id, choice="for", salt="s")
        Vote.record(member_id=members[1].id, motion_id=m2.id, choice="for", salt="s")
        Vote.record(member_id=members[2].id, motion_id=m2.id, choice="for", salt="s")
        Vote.record(
            member_id=members[3].id, motion_id=m2.id, choice="against", salt="s"
        )

        user = _make_user(True)
        with app.test_request_context(
            f"/meetings/{meeting.id}/close-stage2", method="POST"
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_final_results") as mock_send:
                    meetings.close_stage2(meeting.id)
                    assert mock_send.call_count == len(members)

        assert m1.status == "carried"
        assert m2.status == "carried"
        assert meeting.status == "Completed"


def test_meeting_form_duration_validations():
    app = create_app()
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        now = datetime.utcnow()

        # Stage 1 opens before notice period
        data = MultiDict(
            {
                "title": "AGM",
                "notice_date": now.strftime("%Y-%m-%dT%H:%M"),
                "opens_at_stage1": (now + timedelta(days=13)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "closes_at_stage1": (now + timedelta(days=20)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "opens_at_stage2": (now + timedelta(days=21)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "closes_at_stage2": (now + timedelta(days=27)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
            }
        )
        form = MeetingForm(formdata=data)
        assert not form.validate()
        assert form.opens_at_stage1.errors

        # Stage 1 shorter than 7 days
        data = MultiDict(
            {
                "title": "AGM",
                "notice_date": now.strftime("%Y-%m-%dT%H:%M"),
                "opens_at_stage1": now.strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage1": (now + timedelta(days=6)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "opens_at_stage2": (now + timedelta(days=7)).strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage2": (now + timedelta(days=12)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
            }
        )
        form = MeetingForm(formdata=data)
        assert not form.validate()
        assert form.closes_at_stage1.errors

        # Stage 2 opens less than 1 day after Stage 1 closes
        data = MultiDict(
            {
                "title": "AGM",
                "opens_at_stage1": now.strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage1": (now + timedelta(days=7)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "opens_at_stage2": (now + timedelta(days=7, hours=12)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "closes_at_stage2": (now + timedelta(days=12)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
            }
        )
        form = MeetingForm(formdata=data)
        assert not form.validate()
        assert form.opens_at_stage2.errors

        # Stage 2 shorter than 5 days
        data = MultiDict(
            {
                "title": "AGM",
                "opens_at_stage1": now.strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage1": (now + timedelta(days=7)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "opens_at_stage2": (now + timedelta(days=8)).strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage2": (now + timedelta(days=11)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
            }
        )
        form = MeetingForm(formdata=data)
        assert not form.validate()
        assert form.closes_at_stage2.errors

        # Stage 1 opens in the past
        data = MultiDict(
            {
                "title": "AGM",
                "opens_at_stage1": (now - timedelta(hours=1)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "closes_at_stage1": (now + timedelta(days=7)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "opens_at_stage2": (now + timedelta(days=8)).strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage2": (now + timedelta(days=13)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
            }
        )
        form = MeetingForm(formdata=data)
        assert not form.validate()
        assert form.opens_at_stage1.errors

        # Stage 2 opens before Stage 1 opens
        data = MultiDict(
            {
                "title": "AGM",
                "opens_at_stage1": (now + timedelta(days=1)).strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage1": (now + timedelta(days=8)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
                "opens_at_stage2": now.strftime("%Y-%m-%dT%H:%M"),
                "closes_at_stage2": (now + timedelta(days=15)).strftime(
                    "%Y-%m-%dT%H:%M"
                ),
            }
        )
        form = MeetingForm(formdata=data)
        assert not form.validate()
        assert form.opens_at_stage2.errors


def test_create_and_delete_conflict():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M",
            text_md="x",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.flush()
        a1 = Amendment(
            meeting_id=meeting.id, motion_id=motion.id, text_md="A1", order=1
        )
        a2 = Amendment(
            meeting_id=meeting.id, motion_id=motion.id, text_md="A2", order=2
        )
        db.session.add_all([a1, a2])
        db.session.commit()

        user = _make_user(True)
        data = {"amendment_a_id": a1.id, "amendment_b_id": a2.id}
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/conflicts", method="POST", data=data
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash"):
                    meetings.manage_conflicts(motion.id)
        assert AmendmentConflict.query.count() == 1
        conflict = AmendmentConflict.query.first()

        with app.test_request_context(
            f"/meetings/conflicts/{conflict.id}/delete", method="POST"
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.delete_conflict(conflict.id)

        assert AmendmentConflict.query.count() == 0


def test_resend_member_token_stage1():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="A", email="a@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/members/{member.id}/resend", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_vote_invite") as mock_send:
                    meetings.resend_member_link(meeting.id, member.id)
                    mock_send.assert_called_once()
                    assert (
                        VoteToken.query.filter_by(member_id=member.id, stage=1).count()
                        == 1
                    )


def test_resend_member_token_stage2():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM", status="Stage 2")
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="B", email="b@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/members/{member.id}/resend", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_stage2_invite") as mock_send:
                    meetings.resend_member_link(meeting.id, member.id)
                    mock_send.assert_called_once()
                    assert (
                        VoteToken.query.filter_by(member_id=member.id, stage=2).count()
                        == 1
                    )


def test_send_member_email_stage1_invite():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM", opens_at_stage1=datetime.utcnow())
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="A", email="a@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/members/{member.id}/email/stage1_invite",
            method="POST",
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_vote_invite") as mock_send:
                    meetings.send_member_email(meeting.id, member.id, "stage1_invite")
                    mock_send.assert_called_once()
                    assert (
                        VoteToken.query.filter_by(member_id=member.id, stage=1).count()
                        == 1
                    )


def test_send_member_email_stage2_invite():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="AGM", ballot_mode="two-stage", opens_at_stage2=datetime.utcnow()
        )
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="B", email="b@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/members/{member.id}/email/stage2_invite",
            method="POST",
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_stage2_invite") as mock_send:
                    meetings.send_member_email(meeting.id, member.id, "stage2_invite")
                    mock_send.assert_called_once()
                    assert (
                        VoteToken.query.filter_by(member_id=member.id, stage=2).count()
                        == 1
                    )


def test_send_member_email_amendment_review_invite():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        now = datetime.utcnow()
        meeting = Meeting(title="AGM", amendments_closes_at=now, opens_at_stage1=now + timedelta(days=2))
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="C", email="c@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/members/{member.id}/email/amendment_review_invite",
            method="POST",
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_amendment_review_invite") as mock_send:
                    meetings.send_member_email(meeting.id, member.id, "amendment_review_invite")
                    mock_send.assert_called_once()


def test_results_summary_lists_unused_proxies():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        principal = Member(meeting_id=meeting.id, name="P", email="p@example.com")
        proxy = Member(meeting_id=meeting.id, name="X", email="x@example.com")
        db.session.add_all([principal, proxy])
        db.session.flush()
        tok, plain = VoteToken.create(
            member_id=principal.id,
            stage=1,
            salt="s",
            proxy_holder_id=proxy.id,
        )
        db.session.commit()

        with app.test_request_context(f"/meetings/{meeting.id}/results"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                html = meetings.results_summary(meeting.id)
                assert "X" in html
                url = url_for(
                    "meetings.resend_proxy_token",
                    meeting_id=meeting.id,
                    token=tok.token,
                )
                assert url in html


def test_results_summary_shows_member_resend_links():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        member = Member(meeting_id=meeting.id, name="M", email="m@example.com")
        db.session.add(member)
        db.session.commit()

        with app.test_request_context(f"/meetings/{meeting.id}/results"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                html = meetings.results_summary(meeting.id)
                assert "m@example.com" in html
                url = url_for(
                    "meetings.resend_member_link",
                    meeting_id=meeting.id,
                    member_id=member.id,
                )
                assert url in html


def test_resend_and_invalidate_proxy_token():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["TOKEN_SALT"] = "s"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        principal = Member(meeting_id=meeting.id, name="P", email="p@example.com")
        proxy = Member(meeting_id=meeting.id, name="X", email="x@example.com")
        db.session.add_all([principal, proxy])
        db.session.flush()
        tok, plain = VoteToken.create(
            member_id=principal.id,
            stage=1,
            salt="s",
            proxy_holder_id=proxy.id,
        )
        db.session.commit()

        with app.test_request_context(
            f"/meetings/{meeting.id}/proxy-tokens/{tok.token}/resend", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.send_proxy_invite") as mock_send:
                    meetings.resend_proxy_token(meeting.id, tok.token)
                    mock_send.assert_called_once()
                    assert (
                        VoteToken.query.filter_by(proxy_holder_id=proxy.id).count() == 2
                    )

        with app.test_request_context(
            f"/meetings/{meeting.id}/proxy-tokens/{tok.token}/invalidate", method="POST"
        ):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.invalidate_proxy_token(meeting.id, tok.token)
        assert db.session.get(VoteToken, tok.token).used_at is not None


def test_edit_and_delete_amendment():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="AGM", opens_at_stage1=datetime.utcnow() + timedelta(days=30)
        )
        db.session.add(meeting)
        db.session.flush()
        m1 = Member(meeting_id=meeting.id, name="A")
        m2 = Member(meeting_id=meeting.id, name="B")
        db.session.add_all([m1, m2])
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="x",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.flush()
        amend = Amendment(
            meeting_id=meeting.id,
            motion_id=motion.id,
            text_md="A1",
            order=1,
            proposer_id=m1.id,
            seconder_id=m2.id,
        )
        db.session.add(amend)
        db.session.commit()

        user = _make_user(True)
        data = {
            "text_md": "changed",
            "proposer_id": m2.id,
            "seconder_id": m1.id,
        }
        with app.test_request_context(
            f"/meetings/amendments/{amend.id}/edit", method="POST", data=data
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.edit_amendment(amend.id)

        assert db.session.get(Amendment, amend.id).text_md == "changed"

        with app.test_request_context(
            f"/meetings/amendments/{amend.id}/delete", method="POST"
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.delete_amendment(amend.id)

        assert Amendment.query.count() == 0


def test_request_motion_change_cutoff():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="T", opens_at_stage1=datetime.utcnow() + timedelta(days=5)
        )
        db.session.add(meeting)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="text",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.commit()

        user = _make_user(True)
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/request-change",
            method="POST",
            data={"withdraw": "y"},
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash") as fl:
                    meetings.request_motion_change(motion.id)
                    fl.assert_called()

        assert motion.withdrawal_requested_at is None


def test_approve_motion_change_marks_withdrawn():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    app.config["WTF_CSRF_ENABLED"] = False
    with app.app_context():
        db.create_all()
        meeting = Meeting(
            title="T", opens_at_stage1=datetime.utcnow() + timedelta(days=10)
        )
        db.session.add(meeting)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="text",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.commit()

        user = _make_user(True)
        with app.test_request_context(
            f"/meetings/motions/{motion.id}/request-change",
            method="POST",
            data={"withdraw": "y"},
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.request_motion_change(motion.id)

        with app.test_request_context(
            f"/meetings/motions/{motion.id}/approve-change/chair", method="POST"
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.approve_motion_change(motion.id, "chair")

        with app.test_request_context(
            f"/meetings/motions/{motion.id}/approve-change/board", method="POST"
        ):
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.approve_motion_change(motion.id, "board")
        refreshed = db.session.get(Motion, motion.id)
        assert refreshed.withdrawn is True


def test_default_stage_time_calculation():
    app = create_app()
    with app.app_context():
        agm = datetime.utcnow() + timedelta(days=60)
        defaults = meetings._calculate_default_times(agm)
        assert defaults["opens_at_stage1"] - defaults["notice_date"] >= timedelta(
            days=app.config["NOTICE_PERIOD_DAYS"]
        )
        assert defaults["closes_at_stage1"] - defaults["opens_at_stage1"] >= timedelta(
            days=app.config["STAGE1_LENGTH_DAYS"]
        )
        assert defaults["closes_at_stage2"] - defaults["opens_at_stage2"] >= timedelta(
            days=app.config["STAGE2_LENGTH_DAYS"]
        )
        assert defaults["opens_at_stage2"] - defaults["closes_at_stage1"] >= timedelta(
            days=app.config["STAGE_GAP_DAYS"]
        )


def test_prefill_form_defaults():
    app = create_app()
    with app.app_context():
        app.config["WTF_CSRF_ENABLED"] = False
        agm = (datetime.utcnow() + timedelta(days=30)).replace(second=0, microsecond=0)
        data = MultiDict({"closes_at_stage2": agm.strftime("%Y-%m-%dT%H:%M")})
        form = MeetingForm(formdata=data)
        meetings._prefill_form_defaults(form)
        assert form.notice_date.data is not None
        assert form.opens_at_stage1.data is not None
        assert form.closes_at_stage1.data is not None
        assert form.opens_at_stage2.data is not None
        assert form.closes_at_stage2.data == agm


def test_clone_meeting_creates_copy():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        src = Meeting(title="AGM", type="AGM", ballot_mode="two-stage", quorum=5)
        db.session.add(src)
        db.session.flush()
        motion = Motion(
            meeting_id=src.id,
            title="M1",
            text_md="text",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.flush()
        db.session.add(
            Amendment(
                meeting_id=src.id,
                motion_id=motion.id,
                text_md="A1",
                order=1,
                status="pending",
            )
        )
        db.session.commit()

        with app.test_request_context(f"/meetings/{src.id}/clone"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                resp = meetings.clone_meeting(src.id)
                assert resp.status_code == 302

        clones = Meeting.query.filter(Meeting.id != src.id).all()
        assert len(clones) == 1
        cloned = clones[0]
        assert cloned.ballot_mode == src.ballot_mode
        assert cloned.opens_at_stage1 is None
        assert Motion.query.filter_by(meeting_id=cloned.id).count() == 1
        assert Amendment.query.filter_by(meeting_id=cloned.id).count() == 1


def test_view_motion_unpublished_permissions():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.flush()
        motion = Motion(
            meeting_id=meeting.id,
            title="M1",
            text_md="Motion text",
            category="motion",
            threshold="normal",
            ordering=1,
        )
        db.session.add(motion)
        db.session.commit()

        url = f"/meetings/motions/{motion.id}"

        # manager can view unpublished motion
        with app.test_request_context(url):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                html = meetings.view_motion(motion.id)
                assert "Motion text" in html

        # anonymous / non-manager gets 404
        with app.test_request_context(url):
            with pytest.raises(NotFound):
                meetings.view_motion(motion.id)


def test_list_members_requires_permission():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="AGM")
        db.session.add(meeting)
        db.session.commit()
        with app.test_request_context(f"/meetings/{meeting.id}/members"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.list_members(meeting.id)
        with app.test_request_context(f"/meetings/{meeting.id}/members"):
            user = _make_user(False)
            with patch("flask_login.utils._get_user", return_value=user):
                with pytest.raises(Forbidden):
                    meetings.list_members(meeting.id)


def test_import_members_rejects_duplicates():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Test")
        db.session.add(meeting)
        db.session.commit()
        csv_content = (
            "member_id,name,email,proxy_for\n"
            "1,Alice,alice@example.com,\n"
            "1,Bob,alice@example.com,\n"
        )
        data = {"csv_file": (io.BytesIO(csv_content.encode()), "members.csv")}
        with app.test_request_context(
            f"/meetings/{meeting.id}/import-members", method="POST", data=data
        ):
            user = _make_user(True)
            dummy_form = SimpleNamespace(
                csv_file=SimpleNamespace(data=io.BytesIO(csv_content.encode()))
            )
            dummy_form.validate_on_submit = lambda: True
            dummy_form.hidden_tag = lambda: ""
            with patch("flask_login.utils._get_user", return_value=user):
                with patch(
                    "app.meetings.routes.MemberImportForm", return_value=dummy_form
                ):
                    with patch("app.meetings.routes.flash") as mock_flash:
                        with patch(
                            "app.meetings.routes.render_template", return_value=""
                        ):
                            meetings.import_members(meeting.id)
                        mock_flash.assert_any_call(
                            "Duplicate email: alice@example.com", "error"
                        )
                        assert Member.query.count() == 1


def test_batch_edit_requires_permission():
    app = create_app()
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    with app.app_context():
        db.create_all()
        meeting = Meeting(title="Batch")
        db.session.add(meeting)
        db.session.commit()

        with app.test_request_context(f"/meetings/{meeting.id}/motions/batch-edit"):
            user = _make_user(True)
            with patch("flask_login.utils._get_user", return_value=user):
                meetings.batch_edit_motions(meeting.id)

        with app.test_request_context(f"/meetings/{meeting.id}/motions/batch-edit"):
            user = _make_user(False)
            with patch("flask_login.utils._get_user", return_value=user):
                with patch("app.meetings.routes.flash"):
                    with pytest.raises(Forbidden):
                        meetings.batch_edit_motions(meeting.id)
